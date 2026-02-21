import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import csv
import os
import json
from datetime import datetime
import threading
import tkinter.font as tkFont

class MistakeBookApp:
    def __init__(self, root):
        self.root = root
        self.root.geometry("1000x700")
        
        # 数据文件路径（包含配置信息）
        self.data_file = "mistakebook_data.csv"
        
        # 初始化配置（从CSV文件中读取）
        self.load_config()
        
        # 初始化多语言字典
        self.translations = {
            'zh': {
                'app_title': '错题本管理系统',
                'menu_file': '文件',
                'menu_collection': '合集',
                'menu_statistics': '统计',
                'menu_settings': '设置',
                'menu_help': '帮助',
                'menu_import_csv': '从CSV导入',
                'menu_import_txt': '从TXT导入',
                'menu_import_excel': '从Excel导入',
                'menu_import_word': '从Word导入',
                'menu_export_all_csv': '导出所有数据为CSV',
                'menu_export_all_txt': '导出所有数据为TXT',
                'menu_export_all_excel': '导出所有数据为Excel',
                'menu_export_all_word': '导出所有数据为Word',
                'menu_export_all_pdf': '导出所有数据为PDF',
                'menu_manage_collection': '管理合集',
                'menu_export_collection': '导出选中合集',
                'menu_refresh_stats': '刷新统计',
                'menu_export_stats_chart': '导出统计图表',
                'menu_export_stats_report_txt': '导出统计报告(TXT)',
                'menu_export_stats_report_pdf': '导出统计报告(PDF)',
                'menu_export_stats_charts_pdf': '导出统计图表(PDF)',
                'menu_theme': '主题切换',
                'menu_font': '字体设置',
                'menu_backup': '数据备份',
                'menu_restore': '数据恢复',
                'menu_about': '关于应用',
                'menu_exit': '退出',
                'tab_mistakes': '📚 错题列表',
                'tab_search': '🔍 搜索',
                'tab_import_export': '📤 导入与导出',
                'tab_statistics': '📊 统计',
                'tab_collection': '📚 合集管理',
                'tab_settings': '⚙️ 设置',
                'search_subject': '📘 按科目:',
                'search_difficulty': '📊 按难度:',
                'search_keyword': '🔑 按关键词:',
                'search_date': '📅 按日期:',
                'search_start_date': '开始日期:',
                'search_end_date': '结束日期:',
                'reset_filters': '🔄 重置所有搜索条件',
                'search_results': '🔍 搜索结果',
                'view_detail': '📖 查看详情',
                'edit_item': '✏️ 编辑',
                'delete_item': '❌ 删除',
                'add_mistake': '➕ 添加错题',
                'edit_mistake': '✏️ 编辑错题',
                'time_label': '📅 时间:',
                'subject_label': '📘 科目:',
                'question_label': '📝 题干:',
                'answer_label': '✅ 正确答案:',
                'attachment_label': '📎 附件:',
                'difficulty_label': '📊 难度:',
                'select_file': '📁 选择文件',
                'save': '💾 保存',
                'cancel': '↩️ 取消',
                'all_subjects': '全部',
                'all_difficulties': '全部',
                'difficulty_simple': '简单',
                'difficulty_medium': '中等',
                'difficulty_hard': '困难',
                'yes': '是',
                'no': '否',
                'ok': '确定',
                'warning': '警告',
                'error': '错误',
                'success': '成功',
                'confirm': '确认',
                'import_success': '成功导入 {count} 条错题！',
                'export_success': '数据已导出到 {path}',
                'data_empty': '未找到有效的错题数据，请检查文件格式。',
                'select_mistake': '请先选择要操作的错题！',
                'fields_required': '题干和正确答案不能为空！',
                'delete_confirm': '确定要删除选中的错题吗？',
                'collection_create': '➕ 创建新合集',
                'collection_name': '合集名称:',
                'collection_desc': '合集描述:',
                'collection_list': '📚 合集列表',
                'collection_problems': '题目数',
                'create_collection': '创建合集',
                'reset_form': '重置',
                'view_problems': '查看题目',
                'export_collection': '导出合集',
                'delete_collection': '删除合集',
                'collection_delete_confirm': '确定要删除合集 \'{name}\' 吗？此操作不可恢复！',
                'stats_info': '📈 统计信息',
                'stats_charts': '📊 统计图表',
                'refresh_stats': '🔄 刷新统计',
                'export_charts': '📤 导出图表',
                'export_report_txt': '📝 导出报告(TXT)',
                'export_report_pdf': '📄 导出报告(PDF)',
                'theme_light': '浅色',
                'theme_dark': '深色',
                'font_settings': '🎨 字体设置',
                'font_family': '🔤 字体:',
                'font_size': '📏 字体大小:',
                'apply': '✅ 应用',
                'data_management': '🗂️ 数据管理',
                'clear_all': '🗑️ 清空所有数据',
                'backup_data': '🔄 备份数据',
                'restore_data': '📥 恢复数据',
                'app_settings': '🔧 应用设置',
                'check_updates': '🔄 检查更新',
                'about_app': 'ℹ️ 关于应用',
                'help_info': '❓ 帮助',
                'data_stats': '📈 数据统计',
                'total_mistakes': '总错题数',
                'clear_confirm': '确定要清空所有错题数据吗？此操作不可恢复！',
                'backup_success': '数据已备份到: {path}',
                'restore_confirm': '确定要从备份文件恢复数据吗？当前数据将被覆盖！',
                'restore_success': '数据已从备份恢复！',
                'help_title': '帮助',
                'help_content': '错题本管理系统使用帮助:\n\n一、系统概述\n错题本管理系统是一个功能全面的学习辅助工具，帮助用户高效管理错题，提供错题录入、分类、检索、分析等功能。\n\n二、功能模块详解\n\n1. 错题列表标签页\n   - 功能：查看、管理所有错题记录\n   - 添加错题：点击"添加"按钮，在弹窗中填写题目信息\n     * 时间：自动填入当前时间，也可手动修改\n     * 科目：选择题目所属科目（语文、数学、英语等）\n     * 题干：输入题目内容\n     * 正确答案：填写正确解题思路或答案\n     * 附件：可关联相关文件（图片、文档等）\n     * 难度：标记题目难度（简单、中等、困难）\n   - 编辑错题：选中错题，点击"编辑"按钮进行修改\n   - 删除错题：选中错题，点击"删除"按钮\n   - 查看详情：点击"详情"按钮查看完整题目信息\n\n2. 搜索标签页\n   - 多条件筛选：支持按科目、难度、关键词、日期范围筛选\n   - 智能搜索：输入关键词可同时搜索题干、答案、科目等字段\n   - 操作功能：对搜索结果可执行查看、编辑、删除等操作\n   - 组合筛选：可同时应用多个筛选条件获得精确结果\n\n3. 导入导出标签页\n   - 导入功能：\n     * CSV格式：支持标准CSV格式数据导入\n     * TXT格式：支持特定格式文本导入\n     * Excel格式：支持XLSX格式导入（需安装pandas）\n     * Word格式：支持DOCX格式导入（需安装python-docx）\n   - 导出功能：\n     * CSV格式：导出为标准CSV文件，便于Excel打开\n     * TXT格式：导出为文本格式，便于阅读\n     * Excel格式：导出为XLSX格式（需安装pandas）\n     * Word格式：导出为DOCX格式（需安装python-docx）\n     * PDF格式：导出为PDF文档（需安装reportlab）\n   - 批量操作：支持导出全部数据或仅选中数据\n\n4. 合集管理标签页（新增）\n   - 创建合集：可创建具有名称和描述的错题合集\n   - 管理合集：查看、编辑、删除已创建的合集\n   - 题目分配：可将错题分配到特定合集\n   - 导出合集：可将指定合集单独导出为文件\n\n5. 统计标签页\n   - 数据概览：显示总体错题数、各科目分布、各难度分布\n   - 可视化图表：以饼图和柱状图形式展示数据分布\n   - 报告生成：可导出统计报告和图表\n   - 实时更新：数据变化时自动更新统计信息\n\n6. 设置标签页\n   - 主题切换：支持浅色/深色主题切换\n   - 字体设置：可调整界面字体大小\n   - 数据管理：提供数据备份与恢复功能\n   - 应用信息：查看版本和关于信息\n\n三、使用技巧\n   1. 建议定期备份数据，防止数据丢失\n   2. 合理使用难度标记，便于后续复习\n   3. 利用搜索功能快速定位特定错题\n   4. 通过统计功能分析学习薄弱环节\n   5. 利用合集功能将相关错题归类整理\n\n四、快捷操作\n   - 选中错题后可直接编辑或查看详情\n   - 可同时选中多条错题进行批量操作\n   - 搜索功能支持实时筛选，输入即响应\n   - 支持拖拽调整窗口大小以获得更好的显示效果\n\n五、常见问题\n   Q: 无法导入Excel文件？\n   A: 请确保已安装pandas库：pip install pandas openpyxl\n\n   Q: 导出PDF失败？\n   A: 请确保已安装reportlab库：pip install reportlab\n\n   Q: 如何备份数据？\n   A: 在设置标签页中点击"数据备份"按钮，选择备份位置即可\n\n   Q: 可以将错题按主题分类吗？\n   A: 可以使用合集功能创建主题合集，将相关错题归类管理',
                'about_content': '错题本管理系统 v1.0\n\n作者: MistakeBook Team\n功能: 管理错题，支持导入导出多种格式\n界面: 使用tkinter构建\n\n感谢使用本系统！'
            },
            'en': {
                'app_title': 'Mistake Book Management System',
                'menu_file': 'File',
                'menu_collection': 'Collection',
                'menu_statistics': 'Statistics',
                'menu_settings': 'Settings',
                'menu_help': 'Help',
                'menu_import_csv': 'Import from CSV',
                'menu_import_txt': 'Import from TXT',
                'menu_import_excel': 'Import from Excel',
                'menu_import_word': 'Import from Word',
                'menu_export_all_csv': 'Export all data to CSV',
                'menu_export_all_txt': 'Export all data to TXT',
                'menu_export_all_excel': 'Export all data to Excel',
                'menu_export_all_word': 'Export all data to Word',
                'menu_export_all_pdf': 'Export all data to PDF',
                'menu_manage_collection': 'Manage Collections',
                'menu_export_collection': 'Export Selected Collection',
                'menu_refresh_stats': 'Refresh Statistics',
                'menu_export_stats_chart': 'Export Statistics Chart',
                'menu_export_stats_report_txt': 'Export Statistics Report (TXT)',
                'menu_export_stats_report_pdf': 'Export Statistics Report (PDF)',
                'menu_export_stats_charts_pdf': 'Export Statistics Charts (PDF)',
                'menu_theme': 'Toggle Theme',
                'menu_font': 'Font Settings',
                'menu_backup': 'Data Backup',
                'menu_restore': 'Data Restore',
                'menu_about': 'About',
                'menu_exit': 'Exit',
                'tab_mistakes': '📚 Mistake List',
                'tab_search': '🔍 Search',
                'tab_import_export': '📤 Import & Export',
                'tab_statistics': '📊 Statistics',
                'tab_collection': '📚 Collection Management',
                'tab_settings': '⚙️ Settings',
                'search_subject': '📘 By Subject:',
                'search_difficulty': '📊 By Difficulty:',
                'search_keyword': '🔑 By Keyword:',
                'search_date': '📅 By Date:',
                'search_start_date': 'Start Date:',
                'search_end_date': 'End Date:',
                'reset_filters': '🔄 Reset All Filters',
                'search_results': '🔍 Search Results',
                'view_detail': '📖 View Detail',
                'edit_item': '✏️ Edit',
                'delete_item': '❌ Delete',
                'add_mistake': '➕ Add Mistake',
                'edit_mistake': '✏️ Edit Mistake',
                'time_label': '📅 Time:',
                'subject_label': '📘 Subject:',
                'question_label': '📝 Question:',
                'answer_label': '✅ Correct Answer:',
                'attachment_label': '📎 Attachment:',
                'difficulty_label': '📊 Difficulty:',
                'select_file': '📁 Select File',
                'save': '💾 Save',
                'cancel': '↩️ Cancel',
                'all_subjects': 'All',
                'all_difficulties': 'All',
                'difficulty_simple': 'Simple',
                'difficulty_medium': 'Medium',
                'difficulty_hard': 'Hard',
                'yes': 'Yes',
                'no': 'No',
                'ok': 'OK',
                'warning': 'Warning',
                'error': 'Error',
                'success': 'Success',
                'confirm': 'Confirm',
                'import_success': 'Successfully imported {count} mistakes!',
                'export_success': 'Data exported to {path}',
                'data_empty': 'No valid mistake data found, please check the file format.',
                'select_mistake': 'Please select a mistake to operate!',
                'fields_required': 'Question and answer cannot be empty!',
                'delete_confirm': 'Are you sure you want to delete the selected mistakes?',
                'collection_create': '➕ Create New Collection',
                'collection_name': 'Collection Name:',
                'collection_desc': 'Collection Description:',
                'collection_list': '📚 Collection List',
                'collection_problems': 'Number of Problems',
                'create_collection': 'Create Collection',
                'reset_form': 'Reset',
                'view_problems': 'View Problems',
                'export_collection': 'Export Collection',
                'delete_collection': 'Delete Collection',
                'collection_delete_confirm': 'Are you sure you want to delete collection \'{name}\'? This action cannot be undone!',
                'stats_info': '📈 Statistics Info',
                'stats_charts': '📊 Statistics Charts',
                'refresh_stats': '🔄 Refresh Statistics',
                'export_charts': '📤 Export Charts',
                'export_report_txt': '📝 Export Report (TXT)',
                'export_report_pdf': '📄 Export Report (PDF)',
                'theme_light': 'Light',
                'theme_dark': 'Dark',
                'font_settings': '🎨 Font Settings',
                'font_family': '🔤 Font:',
                'font_size': '📏 Font Size:',
                'apply': '✅ Apply',
                'data_management': '🗂️ Data Management',
                'clear_all': '🗑️ Clear All Data',
                'backup_data': '🔄 Backup Data',
                'restore_data': '📥 Restore Data',
                'app_settings': '🔧 App Settings',
                'check_updates': '🔄 Check for Updates',
                'about_app': 'ℹ️ About App',
                'help_info': '❓ Help',
                'data_stats': '📈 Data Statistics',
                'total_mistakes': 'Total Mistakes',
                'clear_confirm': 'Are you sure you want to clear all mistake data? This action cannot be undone!',
                'backup_success': 'Data backed up to: {path}',
                'restore_confirm': 'Are you sure you want to restore data from backup? Current data will be overwritten!',
                'restore_success': 'Data restored from backup!',
                'help_title': 'Help',
                'help_content': 'Mistake Book Management System User Guide:\n\n1. System Overview\nThe Mistake Book Management System is a comprehensive learning aid tool that helps users efficiently manage mistakes, providing functions such as mistake entry, classification, retrieval, and analysis.\n\n2. Functional Modules\n\n1. Mistake List Tab\n   - Function: View and manage all mistake records\n   - Add Mistake: Click the "Add" button and fill in the question information in the pop-up window\n     * Time: Automatically filled with current time, can also be manually modified\n     * Subject: Select the subject to which the question belongs (Chinese, Mathematics, English, etc.)\n     * Question: Enter the question content\n     * Correct Answer: Fill in the correct solution idea or answer\n     * Attachment: Can associate related files (images, documents, etc.)\n     * Difficulty: Mark the question difficulty (Simple, Medium, Hard)\n   - Edit Mistake: Select the mistake and click the "Edit" button to modify\n   - Delete Mistake: Select the mistake and click the "Delete" button\n   - View Details: Click the "Details" button to view complete question information\n\n2. Search Tab\n   - Multi-condition filtering: Support filtering by subject, difficulty, keyword, date range\n   - Smart search: Enter keywords to search across fields such as question, answer, subject, etc.\n   - Operation functions: Perform operations such as viewing, editing, deleting on search results\n   - Combined filtering: Apply multiple filtering conditions simultaneously for precise results\n\n3. Import/Export Tab\n   - Import Functions:\n     * CSV format: Support standard CSV format data import\n     * TXT format: Support specific format text import\n     * Excel format: Support XLSX format import (requires pandas)\n     * Word format: Support DOCX format import (requires python-docx)\n   - Export Functions:\n     * CSV format: Export as standard CSV file, suitable for Excel\n     * TXT format: Export as text format, suitable for reading\n     * Excel format: Export as XLSX format (requires pandas)\n     * Word format: Export as DOCX format (requires python-docx)\n     * PDF format: Export as PDF document (requires reportlab)\n   - Batch operations: Support exporting all data or only selected data\n\n4. Collection Management Tab (New)\n   - Create Collection: Create mistake collections with names and descriptions\n   - Manage Collections: View, edit, delete created collections\n   - Problem Assignment: Assign mistakes to specific collections\n   - Export Collections: Export specified collections as separate files\n\n5. Statistics Tab\n   - Data overview: Display total number of mistakes, distribution by subject, distribution by difficulty\n   - Visualization: Display data distribution in pie charts and bar charts\n   - Report generation: Export statistical reports and charts\n   - Real-time updates: Automatically update statistics when data changes\n\n6. Settings Tab\n   - Theme switching: Support light/dark theme switching\n   - Font settings: Adjust interface font size\n   - Data management: Provide data backup and restore functions\n   - App information: View version and about information\n\n3. Usage Tips\n   1. Regularly back up data to prevent data loss\n   2. Use difficulty marking appropriately for subsequent review\n   3. Use search function to quickly locate specific mistakes\n   4. Analyze learning weak points through statistical functions\n   5. Use collection function to categorize related mistakes\n\n4. Quick Operations\n   - Directly edit or view details after selecting mistakes\n   - Select multiple mistakes simultaneously for batch operations\n   - Search function supports real-time filtering, responding as you type\n   - Support dragging to adjust window size for better display\n\n5. Common Questions\n   Q: Cannot import Excel file?\n   A: Please ensure pandas library is installed: pip install pandas openpyxl\n\n   Q: Export to PDF failed?\n   A: Please ensure reportlab library is installed: pip install reportlab\n\n   Q: How to back up data?\n   A: Click the "Data Backup" button in the settings tab and select the backup location\n\n   Q: Can I categorize mistakes by topic?\n   A: Yes, use the collection function to create topic collections and categorize related mistakes',
                'about_content': 'Mistake Book Management System v1.0\n\nAuthor: MistakeBook Team\nFeatures: Manage mistakes, support import/export multiple formats\nInterface: Built with tkinter\n\nThank you for using this system!'
            }
        }
        
        # 初始化语言设置（必须在使用get_text方法之前）
        self.language = self.config.get('language', 'zh')  # 默认为中文
        
        # 设置窗口标题
        self.root.title(self.get_text('app_title'))
        
        # 初始化字体
        self.default_font = tkFont.nametofont("TkDefaultFont")
        self.default_font_size = self.default_font.actual()['size']
        
        # 初始化数据（从CSV文件中读取）
        self.data = []
        self.load_data()
        
        # 初始化合集数据
        self.collections = {}  # 存储合集信息 {collection_name: {'description': str, 'problems': [problem_ids]}}
        
        # 创建界面
        self.create_widgets()
        
        # 设置主题
        self.apply_theme()
        
        # 绑定窗口大小调整事件
        self.root.bind('<Configure>', self.on_window_resize)
    def load_config(self):
        """从CSV文件中加载配置信息，如果没有则创建默认配置"""
        default_config = {
            "theme": "light",  # "light" or "dark"
            "font_size": 12,
            "window_width": 1000,
            "window_height": 700,
            "language": "zh"   # 默认语言：zh(中文) 或 en(英文)
        }
        
        # 尝试从CSV文件中读取配置信息
        config_found = False
        if os.path.exists(self.data_file):
            try:
                with open(self.data_file, 'r', encoding='utf-8-sig') as f:
                    first_line = f.readline()
                    
                # 查找配置信息
                if first_line.startswith('#CONFIG#'):
                    # 解析配置行
                    import json as json_module
                    config_str = first_line[len('#CONFIG#'):].strip()
                    try:
                        self.config = json_module.loads(config_str)
                        config_found = True
                    except:
                        pass
            except:
                pass
        
        # 如果没有找到配置，使用默认配置
        if not config_found:
            self.config = default_config
        
        # 设置语言
        self.language = self.config.get('language', 'zh')
    
    def save_config(self):
        """保存配置到CSV文件"""
        # 读取现有数据（除了配置行）
        existing_data = []
        if os.path.exists(self.data_file):
            with open(self.data_file, 'r', encoding='utf-8-sig') as f:
                lines = f.readlines()
                # 跳过配置行，只保留标题和数据行
                for line in lines:
                    if not line.startswith('#CONFIG#'):
                        existing_data.append(line)
        
        # 重新写入配置和数据
        with open(self.data_file, 'w', encoding='utf-8-sig') as f:
            import json as json_module
            config_str = json_module.dumps(self.config, ensure_ascii=False)
            f.write(f'#CONFIG#{config_str}\n')
            
            # 重新写入数据（包括标题行）
            for line in existing_data:
                f.write(line)
    
    def get_text(self, key):
        """获取翻译文本"""
        return self.translations[self.language].get(key, key)
    
    def switch_language(self, lang):
        """切换语言"""
        if lang in self.translations:
            self.language = lang
            self.config['language'] = lang
            self.save_config()  # 保存语言设置到配置
    
    def load_data(self):
        """从CSV文件加载数据（配置已在load_config中加载）"""
        self.data = []
        if os.path.exists(self.data_file):
            with open(self.data_file, 'r', encoding='utf-8-sig') as f:  # 使用utf-8-sig处理BOM
                content = f.read()
            
            # 按行分割内容
            lines = content.split('\n')
            
            # 跳过配置行，找到数据部分
            data_lines = []
            for line in lines:
                if not line.startswith('#CONFIG#') and line.strip():
                    data_lines.append(line)
            
            # 重新组合数据内容
            if len(data_lines) > 1:  # 确保有表头和至少一行数据
                data_content = '\n'.join(data_lines)
                
                # 使用csv模块解析
                import io
                try:
                    reader = csv.DictReader(io.StringIO(data_content))
                    # 过滤掉空行或无效数据
                    for row in reader:
                        if row and row.get('时间') and row.get('科目'):
                            # 如果数据行没有合集字段，添加默认空值
                            if '合集' not in row:
                                row['合集'] = ''  # 空字符串表示不属于任何合集
                            self.data.append(row)
                except Exception as e:
                    print(f"解析CSV数据时出错: {e}")
                    # 如果解析失败，初始化为空列表
                    self.data = []
    
    def save_data(self):
        """保存数据到CSV文件（包含配置信息）"""
        # 读取当前配置
        current_config = self.config if hasattr(self, 'config') else {
            "theme": "light",
            "font_size": 12,
            "window_width": 1000,
            "window_height": 700
        }
        
        # 将配置和数据写入CSV文件
        with open(self.data_file, 'w', newline='', encoding='utf-8-sig') as f:
            # 写入配置信息作为第一行
            import json as json_module
            config_str = json_module.dumps(current_config, ensure_ascii=False)
            f.write(f'#CONFIG#{config_str}\n')
            
            # 写入表头（添加合集字段）
            f.write('时间,科目,题干,正确答案,附件路径,难度,合集\n')
            
            # 写入数据
            if self.data:
                # 确保每条数据都有合集字段
                for item in self.data:
                    if '合集' not in item:
                        item['合集'] = ''
                
                # 将数据写入临时字符串，然后写入文件
                import io
                import csv
                temp_buffer = io.StringIO()
                fieldnames = ['时间', '科目', '题干', '正确答案', '附件路径', '难度', '合集']
                writer = csv.DictWriter(temp_buffer, fieldnames=fieldnames)
                writer.writerows(self.data)
                
                # 获取临时缓冲区的内容并写入文件
                data_content = temp_buffer.getvalue()
                f.write(data_content)

    def create_widgets(self):
        """创建界面组件"""
        # 创建主框架
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建菜单栏
        self.create_menu(main_frame)
        
        # 创建Notebook（标签页控件）
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        # 创建各个标签页
        self.create_mistake_list_tab()    # 错题列表标签页
        self.create_search_tab()          # 搜索标签页
        self.create_import_export_tab()   # 导入与导出标签页
        self.create_statistics_tab()      # 统计标签页
        self.create_collection_manager_tab()  # 合集管理标签页
        self.create_settings_tab()        # 设置标签页
        
        # 创建添加错题界面（隐藏在单独的窗口中）
        self.create_add_interface(main_frame)

    def create_menu(self, parent):
        """创建菜单栏"""
        menubar = tk.Menu(parent.winfo_toplevel())
        
        # 文件菜单
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label=self.get_text('menu_import_csv'), command=self.import_from_csv)
        file_menu.add_command(label=self.get_text('menu_import_txt'), command=self.import_from_txt)
        file_menu.add_command(label=self.get_text('menu_import_excel'), command=self.import_from_excel)
        file_menu.add_command(label=self.get_text('menu_import_word'), command=self.import_from_word)
        file_menu.add_separator()
        file_menu.add_command(label=self.get_text('menu_export_all_csv'), command=self.export_all_to_csv)
        file_menu.add_command(label=self.get_text('menu_export_all_txt'), command=self.export_all_to_txt)
        file_menu.add_command(label=self.get_text('menu_export_all_excel'), command=self.export_all_to_excel)
        file_menu.add_command(label=self.get_text('menu_export_all_word'), command=self.export_all_to_word)
        file_menu.add_command(label=self.get_text('menu_export_all_pdf'), command=self.export_all_to_pdf)
        file_menu.add_separator()
        file_menu.add_command(label=self.get_text('menu_exit'), command=self.root.quit)
        menubar.add_cascade(label=self.get_text('menu_file'), menu=file_menu)
        
        # 合集菜单
        collection_menu = tk.Menu(menubar, tearoff=0)
        collection_menu.add_command(label=self.get_text('menu_manage_collection'), command=lambda: self.notebook.select(4))  # 合集管理标签页索引
        collection_menu.add_command(label=self.get_text('menu_export_collection'), command=self.export_collection)
        menubar.add_cascade(label=self.get_text('menu_collection'), menu=collection_menu)
        
        # 统计菜单
        stats_menu = tk.Menu(menubar, tearoff=0)
        stats_menu.add_command(label=self.get_text('menu_refresh_stats'), command=self.refresh_statistics)
        stats_menu.add_command(label=self.get_text('menu_export_stats_chart'), command=self.export_statistics_chart)
        stats_menu.add_command(label=self.get_text('menu_export_stats_report_txt'), command=self.export_statistics_report)
        stats_menu.add_command(label=self.get_text('menu_export_stats_report_pdf'), command=self.export_statistics_report_pdf)
        stats_menu.add_command(label=self.get_text('menu_export_stats_charts_pdf'), command=self.export_statistics_charts_pdf)
        menubar.add_cascade(label=self.get_text('menu_statistics'), menu=stats_menu)
        
        # 设置菜单
        settings_menu = tk.Menu(menubar, tearoff=0)
        settings_menu.add_command(label=self.get_text('menu_theme'), command=self.toggle_theme)
        settings_menu.add_command(label=self.get_text('menu_font'), command=self.open_font_settings)
        settings_menu.add_command(label=self.get_text('menu_backup'), command=self.backup_data)
        settings_menu.add_command(label=self.get_text('menu_restore'), command=self.restore_data)
        menubar.add_cascade(label=self.get_text('menu_settings'), menu=settings_menu)
        
        self.root.config(menu=menubar)

    def create_main_interface(self, parent):
        """创建主界面"""
        self.main_frame = ttk.Frame(parent)
        
        # 搜索框
        search_frame = ttk.LabelFrame(self.main_frame, text="🔍 搜索", padding=10)
        search_frame.pack(fill=tk.X, pady=5)
        
        search_container = ttk.Frame(search_frame)
        search_container.pack(fill=tk.X)
        
        self.search_var = tk.StringVar()
        search_entry = ttk.Entry(search_container, textvariable=self.search_var, width=30, font=('Microsoft YaHei', 10))
        search_entry.pack(side=tk.LEFT, padx=(0, 10), fill=tk.X, expand=True)
        search_entry.bind('<KeyRelease>', self.search_data)
        
        ttk.Button(search_container, text="🔍 搜索", command=self.search_data).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(search_container, text="🔄 刷新", command=self.refresh_data).pack(side=tk.LEFT)
        
        # 表格框架
        table_frame = ttk.LabelFrame(self.main_frame, text="📚 错题列表", padding=10)
        table_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 创建表格
        columns = ('时间', '科目', '题干', '正确答案', '附件', '难度')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=12)
        
        # 设置列标题和宽度
        col_widths = {'时间': 140, '科目': 100, '题干': 220, '正确答案': 180, '附件': 80, '难度': 80}
        for col in columns:
            self.tree.heading(col, text=col, anchor=tk.CENTER)
            self.tree.column(col, width=col_widths.get(col, 100), anchor=tk.W)
        
        # 添加滚动条
        v_scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.tree.yview)
        h_scrollbar = ttk.Scrollbar(table_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        # 布局
        self.tree.grid(row=0, column=0, sticky='nsew')
        v_scrollbar.grid(row=0, column=1, sticky='ns')
        h_scrollbar.grid(row=1, column=0, sticky='ew')
        
        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)
        
        # 按钮框架 - 使用更紧凑的布局
        button_frame = ttk.Frame(self.main_frame)
        button_frame.pack(fill=tk.X, pady=5)
        
        # 创建样式
        style = ttk.Style()
        style.configure('Action.TButton', font=('Microsoft YaHei', 10))
        
        # 使用网格布局使按钮排列更整齐
        ttk.Button(button_frame, text="➕ 添加", command=self.show_add_interface, style='Action.TButton').grid(row=0, column=0, padx=2, sticky='ew')
        ttk.Button(button_frame, text="❌ 删除", command=self.delete_selected, style='Action.TButton').grid(row=0, column=1, padx=2, sticky='ew')
        ttk.Button(button_frame, text="✏️ 编辑", command=self.edit_selected, style='Action.TButton').grid(row=0, column=2, padx=2, sticky='ew')
        ttk.Button(button_frame, text="📖 详情", command=self.view_detail, style='Action.TButton').grid(row=0, column=3, padx=2, sticky='ew')
        
        # 配置列权重，使按钮平均分布
        for i in range(4):
            button_frame.grid_columnconfigure(i, weight=1)
        
        # 加载数据到表格
        self.refresh_data()

    def create_mistake_list_tab(self):
        """创建错题列表标签页"""
        # 创建框架
        self.mistake_list_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.mistake_list_frame, text=self.get_text('tab_mistakes'))
        
        # 搜索框
        search_frame = ttk.LabelFrame(self.mistake_list_frame, text="🔍 "+self.get_text('search_keyword').split(':')[0], padding=10)
        search_frame.pack(fill=tk.X, pady=5)
        
        search_container = ttk.Frame(search_frame)
        search_container.pack(fill=tk.X)
        
        self.search_var = tk.StringVar()
        search_entry = ttk.Entry(search_container, textvariable=self.search_var, width=30, font=('Microsoft YaHei', 10))
        search_entry.pack(side=tk.LEFT, padx=(0, 10), fill=tk.X, expand=True)
        search_entry.bind('<KeyRelease>', self.search_data)
        
        ttk.Button(search_container, text="🔍 搜索", command=self.search_data).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(search_container, text="🔄 刷新", command=self.refresh_data).pack(side=tk.LEFT)
        
        # 表格框架
        table_frame = ttk.LabelFrame(self.mistake_list_frame, text="📖 所有错题", padding=10)
        table_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 创建表格
        columns = ('时间', '科目', '题干', '正确答案', '附件', '难度')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=12)
        
        # 设置列标题和宽度
        col_widths = {'时间': 140, '科目': 80, '题干': 350, '正确答案': 300, '附件': 60, '难度': 60}
        for col in columns:
            self.tree.heading(col, text=col, anchor=tk.CENTER)
            self.tree.column(col, width=col_widths.get(col, 100), anchor=tk.W)
        
        # 添加滚动条
        v_scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.tree.yview)
        h_scrollbar = ttk.Scrollbar(table_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        # 布局
        self.tree.grid(row=0, column=0, sticky='nsew')
        v_scrollbar.grid(row=0, column=1, sticky='ns')
        h_scrollbar.grid(row=1, column=0, sticky='ew')
        
        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)
        
        # 按钮框架 - 使用更紧凑的布局
        button_frame = ttk.Frame(self.mistake_list_frame)
        button_frame.pack(fill=tk.X, pady=5)
        
        # 创建样式
        style = ttk.Style()
        style.configure('Action.TButton', font=('Microsoft YaHei', 10))
        
        # 使用网格布局使按钮排列更整齐
        ttk.Button(button_frame, text="➕ 添加", command=self.show_add_interface, style='Action.TButton').grid(row=0, column=0, padx=2, sticky='ew')
        ttk.Button(button_frame, text="❌ 删除", command=self.delete_selected, style='Action.TButton').grid(row=0, column=1, padx=2, sticky='ew')
        ttk.Button(button_frame, text="✏️ 编辑", command=self.edit_selected, style='Action.TButton').grid(row=0, column=2, padx=2, sticky='ew')
        ttk.Button(button_frame, text="📖 详情", command=self.view_detail, style='Action.TButton').grid(row=0, column=3, padx=2, sticky='ew')
        
        # 配置列权重，使按钮平均分布
        for i in range(4):
            button_frame.grid_columnconfigure(i, weight=1)
        
        # 加载数据到表格
        self.refresh_data()

    def create_add_interface(self, parent):
        """创建添加错题界面"""
        self.add_frame = ttk.Frame(parent)
        
        # 表单框架
        form_frame = ttk.LabelFrame(self.add_frame, text="➕ 添加错题", padding=20)
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # 时间
        time_frame = ttk.Frame(form_frame)
        time_frame.grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=5)
        ttk.Label(time_frame, text="📅 时间:", font=('Microsoft YaHei', 10, 'bold')).pack(side=tk.LEFT)
        self.time_var = tk.StringVar(value=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ttk.Entry(time_frame, textvariable=self.time_var, state='readonly', width=25).pack(side=tk.LEFT, padx=(10, 0))
        
        # 科目
        subject_frame = ttk.Frame(form_frame)
        subject_frame.grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=5)
        ttk.Label(subject_frame, text="📘 科目:", font=('Microsoft YaHei', 10, 'bold')).pack(side=tk.LEFT)
        self.subject_var = tk.StringVar()
        subject_combo = ttk.Combobox(subject_frame, textvariable=self.subject_var, 
                                    values=["语文", "数学", "英语", "物理", "化学", "生物", "历史", "地理", "道法"], 
                                    state="readonly", width=23)
        subject_combo.pack(side=tk.LEFT, padx=(10, 0))
        subject_combo.set("语文")
        
        # 题干
        ttk.Label(form_frame, text="📝 题干:", font=('Microsoft YaHei', 10, 'bold')).grid(row=2, column=0, sticky=tk.NW, pady=5)
        self.question_frame = ttk.Frame(form_frame)
        self.question_frame.grid(row=2, column=1, pady=5, padx=(10, 0), sticky='nsew')
        
        self.question_text = tk.Text(self.question_frame, height=6, wrap=tk.WORD, font=('Microsoft YaHei', 10))
        question_scrollbar = ttk.Scrollbar(self.question_frame, orient=tk.VERTICAL, command=self.question_text.yview)
        self.question_text.configure(yscrollcommand=question_scrollbar.set)
        
        self.question_text.grid(row=0, column=0, sticky='nsew')
        question_scrollbar.grid(row=0, column=1, sticky='ns')
        self.question_frame.grid_rowconfigure(0, weight=1)
        self.question_frame.grid_columnconfigure(0, weight=1)
        
        # 正确答案
        ttk.Label(form_frame, text="✅ 正确答案:", font=('Microsoft YaHei', 10, 'bold')).grid(row=3, column=0, sticky=tk.NW, pady=5)
        self.answer_frame = ttk.Frame(form_frame)
        self.answer_frame.grid(row=3, column=1, pady=5, padx=(10, 0), sticky='nsew')
        
        self.answer_text = tk.Text(self.answer_frame, height=6, wrap=tk.WORD, font=('Microsoft YaHei', 10))
        answer_scrollbar = ttk.Scrollbar(self.answer_frame, orient=tk.VERTICAL, command=self.answer_text.yview)
        self.answer_text.configure(yscrollcommand=answer_scrollbar.set)
        
        self.answer_text.grid(row=0, column=0, sticky='nsew')
        answer_scrollbar.grid(row=0, column=1, sticky='ns')
        self.answer_frame.grid_rowconfigure(0, weight=1)
        self.answer_frame.grid_columnconfigure(0, weight=1)
        
        # 附件
        attachment_frame = ttk.Frame(form_frame)
        attachment_frame.grid(row=4, column=0, columnspan=2, sticky=tk.W, pady=5)
        ttk.Label(attachment_frame, text="📎 附件:", font=('Microsoft YaHei', 10, 'bold')).pack(side=tk.LEFT)
        self.attachment_frame = ttk.Frame(attachment_frame)
        self.attachment_frame.pack(side=tk.LEFT, padx=(10, 0))
        
        self.attachment_var = tk.StringVar()
        ttk.Entry(self.attachment_frame, textvariable=self.attachment_var, width=45, state='readonly').grid(row=0, column=0, padx=(0, 5))
        ttk.Button(self.attachment_frame, text="📁 选择文件", command=self.select_attachment).grid(row=0, column=1)
        
        # 难度
        difficulty_frame = ttk.Frame(form_frame)
        difficulty_frame.grid(row=5, column=0, columnspan=2, sticky=tk.W, pady=5)
        ttk.Label(difficulty_frame, text="📊 难度:", font=('Microsoft YaHei', 10, 'bold')).pack(side=tk.LEFT)
        self.difficulty_var = tk.StringVar(value="中等")
        difficulty_combo = ttk.Combobox(difficulty_frame, textvariable=self.difficulty_var, 
                                       values=["简单", "中等", "困难"], state="readonly", width=23)
        difficulty_combo.pack(side=tk.LEFT, padx=(10, 0))
        
        # 按钮框架
        button_frame = ttk.Frame(form_frame)
        button_frame.grid(row=6, column=0, columnspan=2, pady=20)
        
        style = ttk.Style()
        style.configure('Save.TButton', font=('Microsoft YaHei', 10, 'bold'))
        
        ttk.Button(button_frame, text="💾 保存", command=self.save_mistake, style='Save.TButton').pack(side=tk.LEFT, padx=(0, 20))
        ttk.Button(button_frame, text="↩️ 取消", command=self.hide_add_interface).pack(side=tk.LEFT)
        
        # 配置行权重使文本框可以扩展
        form_frame.grid_rowconfigure(2, weight=1)
        form_frame.grid_rowconfigure(3, weight=1)
        form_frame.grid_columnconfigure(1, weight=1)

    def select_attachment(self):
        """选择附件文件"""
        file_path = filedialog.askopenfilename(
            title="选择附件文件",
            filetypes=[
                ("图片文件", "*.jpg *.jpeg *.png *.gif *.bmp"),
                ("文档文件", "*.pdf *.doc *.docx *.txt"),
                ("所有文件", "*.*")
            ]
        )
        if file_path:
            self.attachment_var.set(file_path)

    def save_mistake(self):
        """保存错题"""
        # 获取数据
        time_val = self.time_var.get()
        subject_val = self.subject_var.get()
        question_val = self.question_text.get("1.0", tk.END).strip()
        answer_val = self.answer_text.get("1.0", tk.END).strip()
        attachment_val = self.attachment_var.get()
        difficulty_val = self.difficulty_var.get()
        
        # 验证输入
        if not question_val or not answer_val:
            messagebox.showwarning("警告", "题干和正确答案不能为空！")
            return
        
        # 创建数据字典
        new_data = {
            '时间': time_val,
            '科目': subject_val,
            '题干': question_val,
            '正确答案': answer_val,
            '附件路径': attachment_val,
            '难度': difficulty_val
        }
        
        # 添加到数据列表
        if hasattr(self, 'current_edit_index'):
            # 如果是编辑模式，更新现有数据
            self.data[self.current_edit_index] = new_data
            delattr(self, 'current_edit_index')  # 清除编辑索引
        else:
            # 如果是添加模式，添加新数据
            self.data.append(new_data)
        
        # 保存到文件
        self.save_data()
        
        # 刷新所有界面
        self.refresh_data()  # 刷新错题列表标签页
        self.update_search_results()  # 刷新搜索结果
        self.update_statistics()  # 刷新导入导出标签页统计
        self.update_settings_statistics()  # 刷新设置标签页统计
        
        # 隐藏添加界面
        self.hide_add_interface()
        
        messagebox.showinfo("成功", "错题已保存！")

    def create_search_tab(self):
        """创建搜索标签页"""
        # 创建框架
        self.search_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.search_frame, text=self.get_text('tab_search'))
        
        # 搜索条件框架
        search_condition_frame = ttk.LabelFrame(self.search_frame, text="🔍 "+self.get_text('search_keyword').split(':')[0]+" "+self.get_text('all_subjects'), padding=10)
        search_condition_frame.pack(fill=tk.X, pady=5)
        
        # 多条件搜索框架
        multi_search_frame = ttk.Frame(search_condition_frame)
        multi_search_frame.pack(fill=tk.X, pady=5)
        
        # 按科目搜索
        ttk.Label(multi_search_frame, text=self.get_text('subject_label')[:-1]+":", font=('Microsoft YaHei', 10)).grid(row=0, column=0, sticky=tk.W, padx=(0, 5))
        self.search_subject_var = tk.StringVar()
        subject_combo = ttk.Combobox(multi_search_frame, textvariable=self.search_subject_var, 
                                    values=[self.get_text('all_subjects'), "语文", "数学", "英语", "物理", "化学", "生物", "历史", "地理", "道法"], 
                                    state="readonly", width=12)
        subject_combo.grid(row=0, column=1, padx=(0, 10))
        subject_combo.set(self.get_text('all_subjects'))
        
        # 按难度搜索
        ttk.Label(multi_search_frame, text=self.get_text('difficulty_label')[:-1]+":", font=('Microsoft YaHei', 10)).grid(row=0, column=2, sticky=tk.W, padx=(0, 5))
        self.search_difficulty_var = tk.StringVar()
        difficulty_combo = ttk.Combobox(multi_search_frame, textvariable=self.search_difficulty_var, 
                                       values=[self.get_text('all_difficulties'), self.get_text('difficulty_simple'), self.get_text('difficulty_medium'), self.get_text('difficulty_hard')], state="readonly", width=12)
        difficulty_combo.grid(row=0, column=3, padx=(0, 10))
        difficulty_combo.set(self.get_text('all_difficulties'))
        
        # 按关键词搜索
        ttk.Label(multi_search_frame, text=self.get_text('search_keyword'), font=('Microsoft YaHei', 10)).grid(row=0, column=4, sticky=tk.W, padx=(0, 5))
        self.search_keyword_var = tk.StringVar()
        keyword_entry = ttk.Entry(multi_search_frame, textvariable=self.search_keyword_var, width=20)
        keyword_entry.grid(row=0, column=5, padx=(0, 10), sticky=tk.EW)
        keyword_entry.bind('<KeyRelease>', self.combined_search)
        
        # 按日期搜索
        ttk.Label(multi_search_frame, text=self.get_text('search_date'), font=('Microsoft YaHei', 10)).grid(row=0, column=6, sticky=tk.W, padx=(0, 5))
        ttk.Label(multi_search_frame, text=self.get_text('search_start_date')[:-1], font=('Microsoft YaHei', 10)).grid(row=0, column=7, sticky=tk.W)
        self.search_start_date_var = tk.StringVar()
        start_date_entry = ttk.Entry(multi_search_frame, textvariable=self.search_start_date_var, width=10)
        start_date_entry.grid(row=0, column=8, padx=(0, 5))
        ttk.Label(multi_search_frame, text=self.get_text('search_end_date')[:-1], font=('Microsoft YaHei', 10)).grid(row=0, column=9, sticky=tk.W)
        self.search_end_date_var = tk.StringVar()
        end_date_entry = ttk.Entry(multi_search_frame, textvariable=self.search_end_date_var, width=10)
        end_date_entry.grid(row=0, column=10, padx=(0, 5))
        
        # 配置列权重
        multi_search_frame.columnconfigure(5, weight=1)  # 关键词列可扩展
        
        # 搜索和重置按钮
        button_frame = ttk.Frame(search_condition_frame)
        button_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(button_frame, text="🔍 "+self.get_text('search_keyword').split(':')[0], command=self.combined_search, style='Action.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="🔄 "+self.get_text('reset_filters').split()[1], command=self.reset_search_filters, style='Action.TButton').pack(side=tk.LEFT, padx=5)
        
        # 结果表格框架
        table_frame = ttk.LabelFrame(self.search_frame, text="🔍 "+self.get_text('search_results').split()[1], padding=10)
        table_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 创建表格
        columns = ('时间', '科目', '题干', '正确答案', '附件', '难度')
        self.search_tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=15)
        
        # 设置列标题和宽度
        col_widths = {'时间': 140, '科目': 80, '题干': 350, '正确答案': 300, '附件': 60, '难度': 60}
        for col in columns:
            self.search_tree.heading(col, text=col, anchor=tk.CENTER)
            self.search_tree.column(col, width=col_widths.get(col, 100), anchor=tk.W)
        
        # 添加滚动条
        v_scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.search_tree.yview)
        h_scrollbar = ttk.Scrollbar(table_frame, orient=tk.HORIZONTAL, command=self.search_tree.xview)
        self.search_tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        # 布局
        self.search_tree.grid(row=0, column=0, sticky='nsew')
        v_scrollbar.grid(row=0, column=1, sticky='ns')
        h_scrollbar.grid(row=1, column=0, sticky='ew')
        
        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)
        
        # 操作按钮框架
        action_button_frame = ttk.Frame(self.search_frame)
        action_button_frame.pack(fill=tk.X, pady=5)
        
        style = ttk.Style()
        style.configure('Action.TButton', font=('Microsoft YaHei', 10))
        
        ttk.Button(action_button_frame, text="📖 "+self.get_text('view_detail').split()[1], command=self.view_search_detail, style='Action.TButton').grid(row=0, column=0, padx=2, sticky='ew')
        ttk.Button(action_button_frame, text="✏️ "+self.get_text('edit_item').split()[1], command=self.edit_search_item, style='Action.TButton').grid(row=0, column=1, padx=2, sticky='ew')
        ttk.Button(action_button_frame, text="❌ "+self.get_text('delete_item').split()[1], command=self.delete_search_item, style='Action.TButton').grid(row=0, column=2, padx=2, sticky='ew')
        
        # 配置列权重，使按钮平均分布
        for i in range(3):
            action_button_frame.grid_columnconfigure(i, weight=1)
        
        # 初始化搜索结果表格
        self.update_search_results()

    def combined_search(self, event=None):
        """多条件组合搜索"""
        # 获取所有搜索条件
        subject = self.search_subject_var.get()
        difficulty = self.search_difficulty_var.get()
        keyword = self.search_keyword_var.get()
        start_date = self.search_start_date_var.get()
        end_date = self.search_end_date_var.get()
        
        # 组合条件进行搜索
        filtered_data = self.data
        
        # 按科目筛选
        if subject != "全部":
            filtered_data = [item for item in filtered_data if item['科目'] == subject]
        
        # 按难度筛选
        if difficulty != "全部":
            filtered_data = [item for item in filtered_data if item['难度'] == difficulty]
        
        # 按关键词筛选
        if keyword:
            keyword = keyword.lower()
            filtered_data = [item for item in filtered_data 
                           if keyword in item['题干'].lower() or 
                              keyword in item['正确答案'].lower() or 
                              keyword in item['科目'].lower() or 
                              keyword in item['难度'].lower()]
        
        # 按日期筛选
        if start_date or end_date:
            filtered_data = [item for item in filtered_data 
                           if self.is_date_in_range(item['时间'], start_date, end_date)]
        
        # 更新搜索结果表格
        self.update_search_results_with_data(filtered_data)

    def refresh_data(self):
        """刷新数据显示"""
        # 清空现有数据
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # 添加数据到表格
        for row in self.data:
            # 截断题干和答案以适应表格显示
            question = row['题干'][:50] + '...' if len(row['题干']) > 50 else row['题干']
            answer = row['正确答案'][:30] + '...' if len(row['正确答案']) > 30 else row['正确答案']
            attachment = "有" if row['附件路径'] else "无"
            
            self.tree.insert('', tk.END, values=(
                row['时间'],
                row['科目'],
                question,
                answer,
                attachment,
                row['难度']
            ))

    def filter_by_subject(self):
        """按科目筛选"""
        # 使用多条件搜索方法
        self.combined_search()

    def filter_by_difficulty(self):
        """按难度筛选"""
        # 使用多条件搜索方法
        self.combined_search()

    def filter_by_keyword(self, event=None):
        """按关键词筛选"""
        # 使用多条件搜索方法
        self.combined_search()

    def filter_by_date(self):
        """按日期筛选"""
        # 使用多条件搜索方法
        self.combined_search()

    def reset_search_filters(self):
        """重置所有搜索条件"""
        self.search_subject_var.set("全部")
        self.search_difficulty_var.set("全部")
        self.search_keyword_var.set("")
        self.search_start_date_var.set("")
        self.search_end_date_var.set("")
        self.update_search_results()

    def update_search_results_with_data(self, filtered_data):
        """使用过滤后的数据更新搜索结果表格"""
        # 清空现有数据
        for item in self.search_tree.get_children():
            self.search_tree.delete(item)
        
        # 添加数据到表格
        for row in filtered_data:
            # 截断题干和答案以适应表格显示
            question = row['题干'][:50] + '...' if len(row['题干']) > 50 else row['题干']
            answer = row['正确答案'][:30] + '...' if len(row['正确答案']) > 30 else row['正确答案']
            attachment = "有" if row['附件路径'] else "无"
            
            self.search_tree.insert('', tk.END, values=(
                row['时间'],
                row['科目'],
                question,
                answer,
                attachment,
                row['难度']
            ), tags=(row['时间'], row['科目']))  # 添加标签以便后续识别

    def update_search_results(self, filter_type=None, filter_value=None):
        """更新搜索结果表格"""
        # 清空现有数据
        for item in self.search_tree.get_children():
            self.search_tree.delete(item)
        
        # 筛选数据
        filtered_data = self.data
        
        if filter_type == 'subject' and filter_value != "全部":
            filtered_data = [item for item in filtered_data if item['科目'] == filter_value]
        elif filter_type == 'difficulty' and filter_value != "全部":
            filtered_data = [item for item in filtered_data if item['难度'] == filter_value]
        elif filter_type == 'keyword' and filter_value:
            keyword = filter_value.lower()
            filtered_data = [item for item in filtered_data 
                           if keyword in item['题干'].lower() or 
                              keyword in item['正确答案'].lower() or 
                              keyword in item['科目'].lower() or 
                              keyword in item['难度'].lower()]
        elif filter_type == 'date':
            start_date = filter_value.get('start', '')
            end_date = filter_value.get('end', '')
            if start_date or end_date:
                filtered_data = [item for item in filtered_data 
                               if self.is_date_in_range(item['时间'], start_date, end_date)]
        
        # 添加数据到表格
        for row in filtered_data:
            # 截断题干和答案以适应表格显示
            question = row['题干'][:50] + '...' if len(row['题干']) > 50 else row['题干']
            answer = row['正确答案'][:30] + '...' if len(row['正确答案']) > 30 else row['正确答案']
            attachment = "有" if row['附件路径'] else "无"
            
            self.search_tree.insert('', tk.END, values=(
                row['时间'],
                row['科目'],
                question,
                answer,
                attachment,
                row['难度']
            ), tags=(row['时间'], row['科目']))  # 添加标签以便后续识别

    def is_date_in_range(self, item_date_str, start_date_str, end_date_str):
        """检查日期是否在指定范围内"""
        try:
            item_date = datetime.strptime(item_date_str.split()[0], "%Y-%m-%d")
            if start_date_str:
                start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
                if item_date < start_date:
                    return False
            if end_date_str:
                end_date = datetime.strptime(end_date_str, "%Y-%m-%d")
                if item_date > end_date:
                    return False
            return True
        except:
            return True  # 如果日期格式错误，则包含该条目

    def view_search_detail(self):
        """查看搜索结果的详细信息"""
        selected_items = self.search_tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择要查看的错题！")
            return
        
        # 获取选中行的值
        item = selected_items[0]
        values = self.search_tree.item(item, 'values')
        
        # 从完整数据中查找对应条目
        for row in self.data:
            if row['时间'] == values[0] and row['科目'] == values[1]:
                # 创建详情窗口
                detail_window = tk.Toplevel(self.root)
                detail_window.title("📖 错题详情")
                detail_window.geometry("700x600")
                detail_window.transient(self.root)
                detail_window.grab_set()  # 模态窗口
                
                # 创建主框架
                main_frame = ttk.Frame(detail_window)
                main_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
                
                # 标题
                title_label = tk.Label(main_frame, text="🔍 错题详情", font=('Microsoft YaHei', 16, 'bold'))
                title_label.pack(anchor=tk.W, pady=(0, 15))
                
                # 创建文本框和滚动条
                text_frame = ttk.Frame(main_frame)
                text_frame.pack(fill=tk.BOTH, expand=True)
                
                detail_text = tk.Text(text_frame, wrap=tk.WORD, font=('Microsoft YaHei', 11), 
                                     padx=10, pady=10, spacing1=5, spacing3=5)
                scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=detail_text.yview)
                detail_text.configure(yscrollcommand=scrollbar.set)
                
                # 插入详细内容
                detail_content = f"""📅 时间: {row['时间']}

📘 科目: {row['科目']}

📊 难度: {row['难度']}

📝 题干:
{row['题干']}

✅ 正确答案:
{row['正确答案']}

📎 附件路径: {row['附件路径'] if row['附件路径'] else '无'}
"""
                detail_text.insert(tk.END, detail_content)
                detail_text.config(state=tk.DISABLED)  # 设置为只读
                
                # 布局
                detail_text.grid(row=0, column=0, sticky='nsew')
                scrollbar.grid(row=0, column=1, sticky='ns')
                
                text_frame.grid_rowconfigure(0, weight=1)
                text_frame.grid_columnconfigure(0, weight=1)
                main_frame.grid_rowconfigure(1, weight=1)
                main_frame.grid_columnconfigure(0, weight=1)
                
                break

    def edit_search_item(self):
        """编辑搜索结果中的项目"""
        selected_items = self.search_tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择要编辑的错题！")
            return
        
        # 获取选中行的值
        item = selected_items[0]
        values = self.search_tree.item(item, 'values')
        
        # 在数据列表中查找完整数据
        for i, row in enumerate(self.data):
            if row['时间'] == values[0] and row['科目'] == values[1]:
                # 将数据填充到添加界面的控件中
                self.time_var.set(row['时间'])
                self.subject_var.set(row['科目'])
                self.question_text.delete("1.0", tk.END)
                self.question_text.insert("1.0", row['题干'])
                self.answer_text.delete("1.0", tk.END)
                self.answer_text.insert("1.0", row['正确答案'])
                self.attachment_var.set(row['附件路径'])
                self.difficulty_var.set(row['难度'])
                
                # 记录当前编辑的索引
                self.current_edit_index = i
                
                # 显示添加界面（实际上是编辑界面）
                self.show_add_interface()
                break

    def delete_search_item(self):
        """删除搜索结果中的项目"""
        selected_items = self.search_tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择要删除的错题！")
            return
        
        if messagebox.askyesno("确认", "确定要删除选中的错题吗？"):
            # 获取选中行的值
            item = selected_items[0]
            values = self.search_tree.item(item, 'values')
            
            # 在数据列表中查找并删除对应的数据
            for i, row in enumerate(self.data):
                if row['时间'] == values[0] and row['科目'] == values[1]:
                    del self.data[i]
                    break
            
            # 保存并刷新
            self.save_data()
            self.update_search_results()  # 更新搜索结果
            self.refresh_data()  # 更新主列表

    def search_data(self, event=None):
        """搜索数据"""
        query = self.search_var.get().lower()
        if not query:
            self.refresh_data()
            return
        
        # 清空现有数据
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # 搜索匹配的数据
        for row in self.data:
            if (query in row['时间'].lower() or 
                query in row['科目'].lower() or 
                query in row['题干'].lower() or 
                query in row['正确答案'].lower() or 
                query in row['难度'].lower()):
                
                # 截断题干和答案以适应表格显示
                question = row['题干'][:50] + '...' if len(row['题干']) > 50 else row['题干']
                answer = row['正确答案'][:30] + '...' if len(row['正确答案']) > 30 else row['正确答案']
                attachment = "有" if row['附件路径'] else "无"
                
                self.tree.insert('', tk.END, values=(
                    row['时间'],
                    row['科目'],
                    question,
                    answer,
                    attachment,
                    row['难度']
                ))

    def delete_selected(self):
        """删除选中的错题"""
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择要删除的错题！")
            return
        
        if messagebox.askyesno("确认", "确定要删除选中的错题吗？"):
            for item in selected_items:
                # 获取选中行的值
                values = self.tree.item(item, 'values')
                # 在数据列表中查找并删除对应的数据
                for i, row in enumerate(self.data):
                    if row['时间'] == values[0] and row['科目'] == values[1] and \
                       row['题干'][:50] + '...' if len(row['题干']) > 50 else row['题干'] == values[2]:
                        del self.data[i]
                        break
            
            # 保存并刷新
            self.save_data()
            self.refresh_data()

    def edit_selected(self):
        """编辑选中的错题"""
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择要编辑的错题！")
            return
        
        # 获取选中行的值
        item = selected_items[0]
        values = self.tree.item(item, 'values')
        
        # 在数据列表中查找完整数据
        for i, row in enumerate(self.data):
            if row['时间'] == values[0] and row['科目'] == values[1]:
                # 将数据填充到添加界面的控件中
                self.time_var.set(row['时间'])
                self.subject_var.set(row['科目'])
                self.question_text.delete("1.0", tk.END)
                self.question_text.insert("1.0", row['题干'])
                self.answer_text.delete("1.0", tk.END)
                self.answer_text.insert("1.0", row['正确答案'])
                self.attachment_var.set(row['附件路径'])
                self.difficulty_var.set(row['难度'])
                
                # 记录当前编辑的索引
                self.current_edit_index = i
                
                # 显示添加界面（实际上是编辑界面）
                self.show_add_interface()
                break

    def view_detail(self):
        """查看错题详情"""
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择要查看的错题！")
            return
        
        # 获取选中行的值
        item = selected_items[0]
        values = self.tree.item(item, 'values')
        
        # 在数据列表中查找完整数据
        for row in self.data:
            if row['时间'] == values[0] and row['科目'] == values[1]:
                # 创建详情窗口
                detail_window = tk.Toplevel(self.root)
                detail_window.title("📖 错题详情")
                detail_window.geometry("700x600")
                detail_window.transient(self.root)
                detail_window.grab_set()  # 模态窗口
                
                # 创建主框架
                main_frame = ttk.Frame(detail_window)
                main_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
                
                # 标题
                title_label = tk.Label(main_frame, text="🔍 错题详情", font=('Microsoft YaHei', 16, 'bold'))
                title_label.pack(anchor=tk.W, pady=(0, 15))
                
                # 创建文本框和滚动条
                text_frame = ttk.Frame(main_frame)
                text_frame.pack(fill=tk.BOTH, expand=True)
                
                detail_text = tk.Text(text_frame, wrap=tk.WORD, font=('Microsoft YaHei', 11), 
                                     padx=10, pady=10, spacing1=5, spacing3=5)
                scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=detail_text.yview)
                detail_text.configure(yscrollcommand=scrollbar.set)
                
                # 插入详细内容
                detail_content = f"""📅 时间: {row['时间']}

📘 科目: {row['科目']}

📊 难度: {row['难度']}

📝 题干:
{row['题干']}

✅ 正确答案:
{row['正确答案']}

📎 附件路径: {row['附件路径'] if row['附件路径'] else '无'}
"""
                detail_text.insert(tk.END, detail_content)
                detail_text.config(state=tk.DISABLED)  # 设置为只读
                
                # 布局
                detail_text.grid(row=0, column=0, sticky='nsew')
                scrollbar.grid(row=0, column=1, sticky='ns')
                
                text_frame.grid_rowconfigure(0, weight=1)
                text_frame.grid_columnconfigure(0, weight=1)
                main_frame.grid_rowconfigure(1, weight=1)
                main_frame.grid_columnconfigure(0, weight=1)
                
                break

    def show_main_interface(self):
        """显示主界面"""
        self.add_frame.pack_forget()
        self.main_frame.pack(fill=tk.BOTH, expand=True)

    def show_add_interface(self):
        """显示添加界面"""
        # 由于现在使用标签页，我们需要创建一个新的顶层窗口
        if not hasattr(self, 'add_window') or not self.add_window.winfo_exists():
            self.add_window = tk.Toplevel(self.root)
            self.add_window.title("✏️ 编辑错题" if hasattr(self, 'current_edit_index') else "➕ 添加错题")
            self.add_window.geometry("800x700")
            self.add_window.transient(self.root)  # 设置为临时窗口
            self.add_window.grab_set()  # 模态窗口
            
            # 创建添加界面内容
            self.create_add_interface_content(self.add_window)
        else:
            # 在显示已存在的窗口之前，更新标题
            self.add_window.title("✏️ 编辑错题" if hasattr(self, 'current_edit_index') else "➕ 添加错题")
            self.add_window.lift()  # 将窗口置于最前

    def create_add_interface_content(self, parent):
        """创建添加界面内容"""
        # 表单框架
        form_frame = ttk.LabelFrame(parent, text="➕ 添加错题", padding=20)
        form_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=10)
        
        # 时间
        time_frame = ttk.Frame(form_frame)
        time_frame.grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=5)
        ttk.Label(time_frame, text="📅 时间:", font=('Microsoft YaHei', 10, 'bold')).pack(side=tk.LEFT)
        self.time_var = tk.StringVar(value=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ttk.Entry(time_frame, textvariable=self.time_var, state='readonly', width=25).pack(side=tk.LEFT, padx=(10, 0))
        
        # 科目
        subject_frame = ttk.Frame(form_frame)
        subject_frame.grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=5)
        ttk.Label(subject_frame, text="📘 科目:", font=('Microsoft YaHei', 10, 'bold')).pack(side=tk.LEFT)
        self.subject_var = tk.StringVar()
        subject_combo = ttk.Combobox(subject_frame, textvariable=self.subject_var, 
                                    values=["语文", "数学", "英语", "物理", "化学", "生物", "历史", "地理", "道法"], 
                                    state="readonly", width=23)
        subject_combo.pack(side=tk.LEFT, padx=(10, 0))
        subject_combo.set("语文")
        
        # 题干
        ttk.Label(form_frame, text="📝 题干:", font=('Microsoft YaHei', 10, 'bold')).grid(row=2, column=0, sticky=tk.NW, pady=5)
        self.question_frame = ttk.Frame(form_frame)
        self.question_frame.grid(row=2, column=1, pady=5, padx=(10, 0), sticky='nsew')
        
        self.question_text = tk.Text(self.question_frame, height=6, wrap=tk.WORD, font=('Microsoft YaHei', 10))
        question_scrollbar = ttk.Scrollbar(self.question_frame, orient=tk.VERTICAL, command=self.question_text.yview)
        self.question_text.configure(yscrollcommand=question_scrollbar.set)
        
        self.question_text.grid(row=0, column=0, sticky='nsew')
        question_scrollbar.grid(row=0, column=1, sticky='ns')
        self.question_frame.grid_rowconfigure(0, weight=1)
        self.question_frame.grid_columnconfigure(0, weight=1)
        
        # 正确答案
        ttk.Label(form_frame, text="✅ 正确答案:", font=('Microsoft YaHei', 10, 'bold')).grid(row=3, column=0, sticky=tk.NW, pady=5)
        self.answer_frame = ttk.Frame(form_frame)
        self.answer_frame.grid(row=3, column=1, pady=5, padx=(10, 0), sticky='nsew')
        
        self.answer_text = tk.Text(self.answer_frame, height=6, wrap=tk.WORD, font=('Microsoft YaHei', 10))
        answer_scrollbar = ttk.Scrollbar(self.answer_frame, orient=tk.VERTICAL, command=self.answer_text.yview)
        self.answer_text.configure(yscrollcommand=answer_scrollbar.set)
        
        self.answer_text.grid(row=0, column=0, sticky='nsew')
        answer_scrollbar.grid(row=0, column=1, sticky='ns')
        self.answer_frame.grid_rowconfigure(0, weight=1)
        self.answer_frame.grid_columnconfigure(0, weight=1)
        
        # 附件
        attachment_frame = ttk.Frame(form_frame)
        attachment_frame.grid(row=4, column=0, columnspan=2, sticky=tk.W, pady=5)
        ttk.Label(attachment_frame, text="📎 附件:", font=('Microsoft YaHei', 10, 'bold')).pack(side=tk.LEFT)
        self.attachment_frame = ttk.Frame(attachment_frame)
        self.attachment_frame.pack(side=tk.LEFT, padx=(10, 0))
        
        self.attachment_var = tk.StringVar()
        ttk.Entry(self.attachment_frame, textvariable=self.attachment_var, width=45, state='readonly').grid(row=0, column=0, padx=(0, 5))
        ttk.Button(self.attachment_frame, text="📁 选择文件", command=self.select_attachment).grid(row=0, column=1)
        
        # 难度
        difficulty_frame = ttk.Frame(form_frame)
        difficulty_frame.grid(row=5, column=0, columnspan=2, sticky=tk.W, pady=5)
        ttk.Label(difficulty_frame, text="📊 难度:", font=('Microsoft YaHei', 10, 'bold')).pack(side=tk.LEFT)
        self.difficulty_var = tk.StringVar(value="中等")
        difficulty_combo = ttk.Combobox(difficulty_frame, textvariable=self.difficulty_var, 
                                       values=["简单", "中等", "困难"], state="readonly", width=23)
        difficulty_combo.pack(side=tk.LEFT, padx=(10, 0))
        
        # 按钮框架
        button_frame = ttk.Frame(form_frame)
        button_frame.grid(row=6, column=0, columnspan=2, pady=20)
        
        style = ttk.Style()
        style.configure('Save.TButton', font=('Microsoft YaHei', 10, 'bold'))
        
        ttk.Button(button_frame, text="💾 保存", command=self.save_mistake, style='Save.TButton').pack(side=tk.LEFT, padx=(0, 20))
        ttk.Button(button_frame, text="↩️ 取消", command=self.hide_add_interface, style='Action.TButton').pack(side=tk.LEFT)
        
        # 配置行权重使文本框可以扩展
        form_frame.grid_rowconfigure(2, weight=1)
        form_frame.grid_rowconfigure(3, weight=1)
        form_frame.grid_columnconfigure(1, weight=1)
        
        # 如果是在编辑模式，更新标题和填入数据
        if hasattr(self, 'current_edit_index'):
            parent.title("✏️ 编辑错题")
            form_frame.config(text="✏️ 编辑错题")

    def hide_add_interface(self):
        """隐藏添加界面"""
        if hasattr(self, 'add_window') and self.add_window.winfo_exists():
            self.add_window.destroy()
            
        # 清除编辑索引（如果存在）
        if hasattr(self, 'current_edit_index'):
            delattr(self, 'current_edit_index')

    def show_main_interface(self):
        """显示主界面 - 在标签页结构中，这相当于选择错题列表标签页"""
        self.notebook.select(0)  # 选择第一个标签页（错题列表）

    def toggle_theme(self):
        """切换主题"""
        if self.config['theme'] == 'light':
            self.config['theme'] = 'dark'
        else:
            self.config['theme'] = 'light'
        
        self.apply_theme()
        self.save_config()

    def apply_theme(self):
        """应用主题"""
        theme = self.config['theme']
        
        if theme == 'dark':
            # 暗色主题
            self.root.configure(bg='#2e2e2e')
            style = ttk.Style()
            style.theme_use('clam')
            
            # 配置颜色
            style.configure('TFrame', background='#2e2e2e')
            style.configure('TLabel', background='#2e2e2e', foreground='white')
            style.configure('TButton', background='#4a4a4a', foreground='white')
            style.configure('TEntry', fieldbackground='#555555', foreground='white')
            style.configure('Treeview', background='#3e3e3e', foreground='white', fieldbackground='#3e3e3e')
            style.map('TButton', background=[('active', '#6a6a6a')])
        else:
            # 亮色主题
            self.root.configure(bg='white')
            style = ttk.Style()
            style.theme_use('default')
            
            # 配置颜色
            style.configure('TFrame', background='white')
            style.configure('TLabel', background='white', foreground='black')
            style.configure('TButton', background='#f0f0f0', foreground='black')
            style.configure('TEntry', fieldbackground='white', foreground='black')
            style.configure('Treeview', background='white', foreground='black', fieldbackground='white')
            style.map('TButton', background=[('active', '#e0e0e0')])

    def on_window_resize(self, event):
        """窗口大小调整事件"""
        # 仅在根窗口调整大小时才触发（避免子控件调整大小时触发）
        if event.widget == self.root:
            self.config['window_width'] = event.width
            self.config['window_height'] = event.height
            self.save_config()
            
            # 根据窗口大小调整字体
            self.adjust_font_size()

    def adjust_font_size(self):
        """根据窗口大小调整字体"""
        # 获取当前窗口大小
        width = self.root.winfo_width() or self.config['window_width']
        height = self.root.winfo_height() or self.config['window_height']
        
        # 根据窗口大小动态调整字体大小
        new_font_size = max(10, min(16, int((width + height) / 120)))
        
        if new_font_size != self.config['font_size']:
            self.config['font_size'] = new_font_size
            self.save_config()
            
            # 更新默认字体大小
            self.default_font.configure(size=new_font_size)
            
            # 更新特定控件的字体
            self.update_all_fonts(self.default_font.actual()['family'], new_font_size)

    def update_font_size(self, widget, font_size):
        """递归更新控件及其子控件的字体大小"""
        try:
            # 尝试更新控件的字体
            if hasattr(widget, 'configure'):
                # 获取当前字体配置
                current_font = widget.cget('font')
                
                if isinstance(current_font, str):
                    # 如果是字体名称，创建新字体
                    new_font = tkFont.Font(font=widget.cget('font'))
                    new_font.configure(size=font_size)
                    widget.configure(font=new_font)
                elif isinstance(current_font, tuple):
                    # 如果是字体元组 (family, size, ...)
                    font_family = current_font[0] if current_font else "TkDefaultFont"
                    new_font = tkFont.Font(family=font_family, size=font_size)
                    widget.configure(font=new_font)
                elif hasattr(current_font, 'configure'):
                    # 如果是Font对象，直接配置
                    current_font.configure(size=font_size)
        except tk.TclError:
            # 某些控件可能不支持font属性，忽略错误
            pass
        
        # 递归更新子控件
        for child in widget.winfo_children():
            self.update_font_size(child, font_size)

    def open_font_settings(self):
        """打开字体设置窗口"""
        font_window = tk.Toplevel(self.root)
        font_window.title("🎨 字体设置")
        font_window.geometry("400x200")
        font_window.transient(self.root)  # 设置为临时窗口
        font_window.grab_set()  # 模态窗口
        
        # 主框架
        main_frame = ttk.Frame(font_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        title_label = tk.Label(main_frame, text="🎨 字体设置", font=('Microsoft YaHei', 14, 'bold'))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 20))
        
        # 字体族选择
        font_frame = ttk.Frame(main_frame)
        font_frame.grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=5)
        ttk.Label(font_frame, text="🔤 字体:", font=('Microsoft YaHei', 10, 'bold')).pack(side=tk.LEFT)
        
        font_families = tkFont.families()
        # 选取常用字体，包括中英文
        common_fonts = []
        for f in font_families:
            lower_f = f.lower()
            if any(x in lower_f for x in ['microsoft', 'sim', 'song', 'hei', 'kai', 'deja', 'liberation', 'ubuntu', 'times', 'arial', 'helvetica', 'consolas', 'courier', 'comic', 'calibri', 'cambria', 'georgia', 'tahoma', 'trebuchet', 'verdana']):
                common_fonts.append(f)
        # 添加一些常见的中文字体
        chinese_fonts = ['SimSun', 'SimHei', 'Microsoft YaHei', 'KaiTi', 'FangSong', 'NSimSun', 'Microsoft JhengHei']
        for font in chinese_fonts:
            if font in font_families and font not in common_fonts:
                common_fonts.append(font)
        common_fonts = common_fonts[:30]  # 限制显示数量
        self.font_family_var = tk.StringVar(value=self.default_font.actual()['family'])
        font_combo = ttk.Combobox(font_frame, textvariable=self.font_family_var, values=common_fonts, state="readonly", width=25)
        font_combo.pack(side=tk.LEFT, padx=(10, 0))
        
        # 字体大小选择
        size_frame = ttk.Frame(main_frame)
        size_frame.grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=15)
        ttk.Label(size_frame, text="📏 字体大小:", font=('Microsoft YaHei', 10, 'bold')).pack(side=tk.LEFT)
        
        self.font_size_var = tk.IntVar(value=self.config.get('font_size', self.default_font_size))
        size_spinbox = tk.Spinbox(size_frame, from_=8, to=24, textvariable=self.font_size_var, width=10, font=('Microsoft YaHei', 10))
        size_spinbox.pack(side=tk.LEFT, padx=(10, 0))
        
        # 按钮框架
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=30)
        
        style = ttk.Style()
        style.configure('Apply.TButton', font=('Microsoft YaHei', 10, 'bold'))
        
        ttk.Button(button_frame, text="✅ 应用", command=lambda: self.apply_font_settings(font_window), style='Apply.TButton').pack(side=tk.LEFT, padx=(0, 20))
        ttk.Button(button_frame, text="❌ 取消", command=font_window.destroy).pack(side=tk.LEFT)

    def apply_font_settings(self, window):
        """应用字体设置"""
        new_family = self.font_family_var.get()
        new_size = self.font_size_var.get()
        
        # 更新默认字体
        self.default_font.configure(family=new_family, size=new_size)
        
        # 更新配置
        self.config['font_family'] = new_family
        self.config['font_size'] = new_size
        self.save_config()
        
        # 更新所有控件的字体
        self.update_all_fonts(new_family, new_size)
        
        # 关闭窗口
        window.destroy()

    def update_all_fonts(self, font_family, font_size):
        """更新所有控件的字体"""
        self.update_font_size(self.root, font_size)

    def create_import_export_tab(self):
        """创建导入与导出标签页"""
        # 创建框架
        self.import_export_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.import_export_frame, text="📤 导入与导出")
        
        # 导入区域
        import_frame = ttk.LabelFrame(self.import_export_frame, text="📥 导入数据", padding=15)
        import_frame.pack(fill=tk.X, pady=10)
        
        # CSV导入
        ttk.Button(import_frame, text="📂 从CSV文件导入", command=self.import_from_csv, 
                  style='Action.TButton').pack(fill=tk.X, pady=5)
        
        # TXT导入
        ttk.Button(import_frame, text="📄 从TXT文件导入", command=self.import_from_txt, 
                  style='Action.TButton').pack(fill=tk.X, pady=5)
        
        # Excel导入
        ttk.Button(import_frame, text="📊 从Excel文件导入", command=self.import_from_excel, 
                  style='Action.TButton').pack(fill=tk.X, pady=5)
        
        # Word导入
        ttk.Button(import_frame, text="📝 从Word文档导入", command=self.import_from_word, 
                  style='Action.TButton').pack(fill=tk.X, pady=5)
        
        # 导出区域
        export_frame = ttk.LabelFrame(self.import_export_frame, text="📤 导出数据", padding=15)
        export_frame.pack(fill=tk.X, pady=10)
        
        # CSV导出
        ttk.Button(export_frame, text="📂 导出所有数据到CSV", command=self.export_all_to_csv, 
                  style='Action.TButton').pack(fill=tk.X, pady=5)
        
        # TXT导出
        ttk.Button(export_frame, text="📄 导出所有数据到TXT", command=self.export_all_to_txt, 
                  style='Action.TButton').pack(fill=tk.X, pady=5)
        
        # Excel导出
        ttk.Button(export_frame, text="📊 导出所有数据到Excel", command=self.export_all_to_excel, 
                  style='Action.TButton').pack(fill=tk.X, pady=5)
        
        # Word导出
        ttk.Button(export_frame, text="📝 导出所有数据到Word", command=self.export_all_to_word, 
                  style='Action.TButton').pack(fill=tk.X, pady=5)
        
        # PDF导出
        ttk.Button(export_frame, text="📄 导出所有数据到PDF", command=self.export_all_to_pdf, 
                  style='Action.TButton').pack(fill=tk.X, pady=5)
        
        # 选中项导出
        selected_export_frame = ttk.LabelFrame(self.import_export_frame, text="📤 导出选中项", padding=15)
        selected_export_frame.pack(fill=tk.X, pady=10)
        
        # 选择要导出的错题
        ttk.Label(selected_export_frame, text="请先在错题列表标签页选择要导出的错题", 
                 font=('Microsoft YaHei', 10)).pack(pady=5)
        
        export_selected_frame = ttk.Frame(selected_export_frame)
        export_selected_frame.pack(fill=tk.X)
        
        ttk.Button(export_selected_frame, text="📄 导出选中项到TXT", command=self.export_selected_to_txt, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        ttk.Button(export_selected_frame, text="📂 导出选中项到CSV", command=self.export_selected_to_csv, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        ttk.Button(export_selected_frame, text="📊 导出选中项到Excel", command=self.export_selected_to_excel, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        # 数据统计
        stats_frame = ttk.LabelFrame(self.import_export_frame, text="📈 数据统计", padding=15)
        stats_frame.pack(fill=tk.X, pady=10)
        
        self.stats_text = tk.Text(stats_frame, height=8, wrap=tk.WORD, font=('Microsoft YaHei', 10))
        stats_scrollbar = ttk.Scrollbar(stats_frame, orient=tk.VERTICAL, command=self.stats_text.yview)
        self.stats_text.configure(yscrollcommand=stats_scrollbar.set)
        
        self.stats_text.grid(row=0, column=0, sticky='nsew')
        stats_scrollbar.grid(row=0, column=1, sticky='ns')
        stats_frame.grid_rowconfigure(0, weight=1)
        stats_frame.grid_columnconfigure(0, weight=1)
        
        # 更新统计数据
        self.update_statistics()
        
    def import_from_csv(self):
        """从CSV文件导入数据"""
        # 询问选择CSV文件
        file_path = filedialog.askopenfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
            title="从CSV导入"
        )
        
        if not file_path:
            return
        
        try:
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                reader = csv.DictReader(f)
                imported_data = [row for row in reader]
            
            if imported_data:
                # 添加到现有数据中
                self.data.extend(imported_data)
                
                # 保存到CSV文件
                self.save_data()
                
                # 刷新界面
                self.refresh_data()
                self.update_search_results()
                self.update_statistics()
                
                messagebox.showinfo("成功", f"成功导入 {len(imported_data)} 条错题！")
            else:
                messagebox.showwarning("警告", "未找到有效的错题数据，请检查CSV文件格式。")
                
        except Exception as e:
            messagebox.showerror("错误", f"导入失败: {str(e)}")

    def import_from_excel(self):
        """从Excel文件导入数据"""
        try:
            import pandas as pd
            # 询问选择Excel文件
            file_path = filedialog.askopenfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")],
                title="从Excel导入"
            )
            
            if not file_path:
                return
            
            # 读取Excel文件
            df = pd.read_excel(file_path)
            
            # 转换为字典列表
            imported_data = df.to_dict('records')
            
            if imported_data:
                # 添加到现有数据中
                self.data.extend(imported_data)
                
                # 保存到CSV文件
                self.save_data()
                
                # 刷新界面
                self.refresh_data()
                self.update_search_results()
                self.update_statistics()
                
                messagebox.showinfo("成功", f"成功导入 {len(imported_data)} 条错题！")
            else:
                messagebox.showwarning("警告", "未找到有效的错题数据，请检查Excel文件格式。")
                
        except ImportError:
            messagebox.showerror("错误", "需要安装pandas和openpyxl库，请运行: pip install pandas openpyxl")
        except Exception as e:
            messagebox.showerror("错误", f"导入失败: {str(e)}")

    def import_from_word(self):
        """从Word文档导入数据"""
        try:
            from docx import Document
            # 询问选择Word文件
            file_path = filedialog.askopenfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
                title="从Word导入"
            )
            
            if not file_path:
                return
            
            # 读取Word文档
            doc = Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # 解析文档内容（这里使用类似TXT解析的逻辑，可根据实际需求调整）
            imported_data = self.parse_txt_content(text)
            
            if imported_data:
                # 添加到现有数据中
                self.data.extend(imported_data)
                
                # 保存到CSV文件
                self.save_data()
                
                # 刷新界面
                self.refresh_data()
                self.update_search_results()
                self.update_statistics()
                
                messagebox.showinfo("成功", f"成功导入 {len(imported_data)} 条错题！")
            else:
                messagebox.showwarning("警告", "未找到有效的错题数据，请检查Word文档格式。")
                
        except ImportError:
            messagebox.showerror("错误", "需要安装python-docx库，请运行: pip install python-docx")
        except Exception as e:
            messagebox.showerror("错误", f"导入失败: {str(e)}")

    def export_all_to_csv(self):
        """导出所有数据到CSV"""
        # 询问保存位置
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
            title="导出为CSV"
        )
        
        if not file_path:
            return
        
        try:
            with open(file_path, 'w', newline='', encoding='utf-8-sig') as f:
                fieldnames = ['时间', '科目', '题干', '正确答案', '附件路径', '难度']
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(self.data)
            
            messagebox.showinfo("成功", f"数据已导出到 {file_path}")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {str(e)}")

    def export_all_to_txt(self):
        """导出所有数据到TXT"""
        # 询问保存位置
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            title="导出为TXT"
        )
        
        if not file_path:
            return
        
        # 写入数据到TXT文件，使用UTF-8编码
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("错题本导出\n")
            f.write("="*50 + "\n\n")
            
            for i, row in enumerate(self.data, 1):
                f.write(f"第{i}题\n")
                f.write(f"时间: {row['时间']}\n")
                f.write(f"科目: {row['科目']}\n")
                f.write(f"难度: {row['难度']}\n")
                f.write(f"题干: {row['题干']}\n")
                f.write(f"正确答案: {row['正确答案']}\n")
                f.write(f"附件路径: {row['附件路径'] if row['附件路径'] else '无'}\n")
                f.write("-" * 50 + "\n\n")
        
        messagebox.showinfo("成功", f"数据已导出到 {file_path}")

    def export_all_to_excel(self):
        """导出所有数据到Excel"""
        try:
            import pandas as pd
            # 询问保存位置
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                title="导出为Excel"
            )
            
            if not file_path:
                return
            
            # 创建DataFrame
            df = pd.DataFrame(self.data)
            
            # 写入Excel文件
            df.to_excel(file_path, index=False, engine='openpyxl')
            
            messagebox.showinfo("成功", f"数据已导出到 {file_path}")
        except ImportError:
            messagebox.showerror("错误", "需要安装pandas和openpyxl库，请运行: pip install pandas openpyxl")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {str(e)}")

    def export_all_to_word(self):
        """导出所有数据到Word"""
        try:
            from docx import Document
            from docx.shared import Inches
            # 询问保存位置
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
                title="导出为Word"
            )
            
            if not file_path:
                return
            
            # 创建文档
            doc = Document()
            doc.add_heading('错题本导出', 0)
            
            for i, row in enumerate(self.data, 1):
                doc.add_heading(f'第{i}题', level=1)
                doc.add_paragraph(f"时间: {row['时间']}")
                doc.add_paragraph(f"科目: {row['科目']}")
                doc.add_paragraph(f"难度: {row['难度']}")
                doc.add_paragraph(f"题干: {row['题干']}")
                doc.add_paragraph(f"正确答案: {row['正确答案']}")
                doc.add_paragraph(f"附件路径: {row['附件路径'] if row['附件路径'] else '无'}")
                doc.add_paragraph("")  # 空行分隔
            
            # 保存文档
            doc.save(file_path)
            
            messagebox.showinfo("成功", f"数据已导出到 {file_path}")
        except ImportError:
            messagebox.showerror("错误", "需要安装python-docx库，请运行: pip install python-docx")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {str(e)}")

    def export_all_to_pdf(self):
        """导出所有数据到PDF"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import os
            
            # 注册中文字体
            # 首先尝试使用系统字体
            font_name = 'SimSun'
            try:
                # 尝试注册字体
                pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))
            except:
                # 如果SimSun不可用，尝试其他常见中文字体
                try:
                    # 在Windows系统上查找常见字体文件
                    font_paths = [
                        r'C:\Windows\Fonts\simsun.ttc',    # 宋体
                        r'C:\Windows\Fonts\msyh.ttc',      # 微软雅黑
                        r'C:\Windows\Fonts\msyhbd.ttc',    # 微软雅黑粗体
                        r'C:\Windows\Fonts\simhei.ttf',    # 黑体
                    ]
                    
                    font_found = False
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            if 'simsun' in font_path.lower():
                                font_name = 'SimSun'
                            elif 'msyh' in font_path.lower():
                                font_name = 'MicrosoftYaHei'
                            elif 'simhei' in font_path.lower():
                                font_name = 'SimHei'
                            
                            pdfmetrics.registerFont(TTFont(font_name, font_path))
                            font_found = True
                            break
                    
                    if not font_found:
                        # 如果找不到系统字体，则报错
                        messagebox.showerror("错误", "未找到支持中文的字体文件")
                        return
                except:
                    messagebox.showerror("错误", "无法注册中文字体，请确保系统中有中文字体文件")
                    return
            
            # 询问保存位置
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
                title="导出为PDF"
            )
            
            if not file_path:
                return
            
            # 创建PDF文档
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            story = []
            
            # 标题样式（使用中文字体）
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=getSampleStyleSheet()['Title'],
                fontName=font_name,
                fontSize=20,
                spaceAfter=30,
                alignment=1  # 居中
            )
            
            # 正文样式（使用中文字体）
            content_style = ParagraphStyle(
                'CustomContent',
                parent=getSampleStyleSheet()['Normal'],
                fontName=font_name,
                fontSize=12,
                spaceAfter=12
            )
            
            # 添加标题
            title = Paragraph("错题本导出", title_style)
            story.append(title)
            story.append(Spacer(1, 0.2 * inch))
            
            # 添加数据
            for i, row in enumerate(self.data, 1):
                story.append(Paragraph(f"第{i}题", content_style))
                story.append(Paragraph(f"时间: {row['时间']}", content_style))
                story.append(Paragraph(f"科目: {row['科目']}", content_style))
                story.append(Paragraph(f"难度: {row['难度']}", content_style))
                story.append(Paragraph(f"题干: {row['题干']}", content_style))
                story.append(Paragraph(f"正确答案: {row['正确答案']}", content_style))
                story.append(Spacer(1, 0.2 * inch))
            
            # 生成PDF
            doc.build(story)
            messagebox.showinfo("成功", f"数据已导出到 {file_path}")
            
        except ImportError:
            messagebox.showerror("错误", "需要安装reportlab库，请运行: pip install reportlab")
        except Exception as e:
            messagebox.showerror("错误", f"PDF导出失败: {str(e)}")

    def export_selected_to_txt(self):
        """导出选中的错题到TXT"""
        # 获取主错题列表中的选中项
        selected_items = self.tree.selection()  # 使用主列表的tree
        if not selected_items:
            messagebox.showwarning("警告", "请先在错题列表标签页选择要导出的错题！")
            # 切换到错题列表标签页
            self.notebook.select(0)  # 选择第一个标签页
            return
        
        # 获取选中项的数据
        selected_data = []
        for item in selected_items:
            values = self.tree.item(item, 'values')
            # 根据时间和科目找到完整数据
            for row in self.data:
                if row['时间'] == values[0] and row['科目'] == values[1]:
                    selected_data.append(row)
                    break
        
        if not selected_data:
            messagebox.showwarning("警告", "未找到选中的错题数据！")
            return
        
        # 询问保存位置
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            title="导出选中的错题到TXT"
        )
        
        if not file_path:
            return
        
        # 写入选中数据到TXT文件
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("错题本选中项导出\n")
            f.write("="*50 + "\n\n")
            
            for i, row in enumerate(selected_data, 1):
                f.write(f"第{i}题\n")
                f.write(f"时间: {row['时间']}\n")
                f.write(f"科目: {row['科目']}\n")
                f.write(f"难度: {row['难度']}\n")
                f.write(f"题干: {row['题干']}\n")
                f.write(f"正确答案: {row['正确答案']}\n")
                f.write(f"附件路径: {row['附件路径'] if row['附件路径'] else '无'}\n")
                f.write("-" * 50 + "\n\n")
        
        messagebox.showinfo("成功", f"已导出 {len(selected_data)} 条选中的错题到 {file_path}")

    def export_selected_to_csv(self):
        """导出选中的错题到CSV"""
        # 获取主错题列表中的选中项
        selected_items = self.tree.selection()  # 使用主列表的tree
        if not selected_items:
            messagebox.showwarning("警告", "请先在错题列表标签页选择要导出的错题！")
            # 切换到错题列表标签页
            self.notebook.select(0)  # 选择第一个标签页
            return
        
        # 获取选中项的数据
        selected_data = []
        for item in selected_items:
            values = self.tree.item(item, 'values')
            # 根据时间和科目找到完整数据
            for row in self.data:
                if row['时间'] == values[0] and row['科目'] == values[1]:
                    selected_data.append(row)
                    break
        
        if not selected_data:
            messagebox.showwarning("警告", "未找到选中的错题数据！")
            return
        
        # 询问保存位置
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
            title="导出选中的错题到CSV"
        )
        
        if not file_path:
            return
        
        try:
            with open(file_path, 'w', newline='', encoding='utf-8-sig') as f:
                fieldnames = ['时间', '科目', '题干', '正确答案', '附件路径', '难度']
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(selected_data)
            
            messagebox.showinfo("成功", f"已导出 {len(selected_data)} 条选中的错题到 {file_path}")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {str(e)}")

    def export_selected_to_excel(self):
        """导出选中的错题到Excel"""
        try:
            import pandas as pd
            # 获取主错题列表中的选中项
            selected_items = self.tree.selection()  # 使用主列表的tree
            if not selected_items:
                messagebox.showwarning("警告", "请先在错题列表标签页选择要导出的错题！")
                # 切换到错题列表标签页
                self.notebook.select(0)  # 选择第一个标签页
                return
            
            # 获取选中项的数据
            selected_data = []
            for item in selected_items:
                values = self.tree.item(item, 'values')
                # 根据时间和科目找到完整数据
                for row in self.data:
                    if row['时间'] == values[0] and row['科目'] == values[1]:
                        selected_data.append(row)
                        break
            
            if not selected_data:
                messagebox.showwarning("警告", "未找到选中的错题数据！")
                return
            
            # 询问保存位置
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                title="导出选中的错题到Excel"
            )
            
            if not file_path:
                return
            
            # 创建DataFrame并写入Excel
            df = pd.DataFrame(selected_data)
            df.to_excel(file_path, index=False, engine='openpyxl')
            
            messagebox.showinfo("成功", f"已导出 {len(selected_data)} 条选中的错题到 {file_path}")
        except ImportError:
            messagebox.showerror("错误", "需要安装pandas和openpyxl库，请运行: pip install pandas openpyxl")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {str(e)}")

    def create_settings_tab(self):
        """创建设置标签页"""
        # 创建框架
        self.settings_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.settings_frame, text="⚙️ 设置")
        
        # 主题设置
        theme_frame = ttk.LabelFrame(self.settings_frame, text="🎨 主题设置", padding=15)
        theme_frame.pack(fill=tk.X, pady=10)
        
        theme_button_frame = ttk.Frame(theme_frame)
        theme_button_frame.pack(fill=tk.X)
        
        ttk.Button(theme_button_frame, text="🌙 切换深色/浅色主题", command=self.toggle_theme, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=5)
        
        # 字体设置
        font_frame = ttk.LabelFrame(self.settings_frame, text="🔤 字体设置", padding=15)
        font_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(font_frame, text="📏 调整字体大小", command=self.open_font_settings, 
                  style='Action.TButton').pack(fill=tk.X, pady=5)
        
        # 数据管理
        data_frame = ttk.LabelFrame(self.settings_frame, text="🗂️ 数据管理", padding=15)
        data_frame.pack(fill=tk.X, pady=10)
        
        data_button_frame = ttk.Frame(data_frame)
        data_button_frame.pack(fill=tk.X)
        
        ttk.Button(data_button_frame, text="🗑️ 清空所有数据", command=self.clear_all_data, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        ttk.Button(data_button_frame, text="🔄 备份数据", command=self.backup_data, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        ttk.Button(data_button_frame, text="📥 恢复数据", command=self.restore_data, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        # 应用设置
        app_frame = ttk.LabelFrame(self.settings_frame, text="🔧 应用设置", padding=15)
        app_frame.pack(fill=tk.X, pady=10)
        
        app_button_frame = ttk.Frame(app_frame)
        app_button_frame.pack(fill=tk.X)
        
        # 语言切换下拉框
        lang_frame = ttk.Frame(app_button_frame)
        lang_frame.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        ttk.Label(lang_frame, text="🌐 语言/Language:").pack(anchor=tk.W)
        self.language_var = tk.StringVar(value=self.language)
        language_combo = ttk.Combobox(lang_frame, textvariable=self.language_var,
                                     values=["zh", "en"], state="readonly", width=10)
        language_combo.pack(pady=5)
        language_combo.bind('<<ComboboxSelected>>', self.change_language)
        
        ttk.Button(app_button_frame, text="ℹ️ "+self.get_text('about_app')[3:], command=self.show_about, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        ttk.Button(app_button_frame, text="❓ "+self.get_text('help_info')[3:], command=self.show_help, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        # 数据统计（复用导入导出标签页的统计功能）
        stats_frame = ttk.LabelFrame(self.settings_frame, text="📈 数据统计", padding=15)
        stats_frame.pack(fill=tk.X, pady=10)
        
        self.settings_stats_text = tk.Text(stats_frame, height=8, wrap=tk.WORD, font=('Microsoft YaHei', 10))
        stats_scrollbar = ttk.Scrollbar(stats_frame, orient=tk.VERTICAL, command=self.settings_stats_text.yview)
        self.settings_stats_text.configure(yscrollcommand=stats_scrollbar.set)
        
        self.settings_stats_text.grid(row=0, column=0, sticky='nsew')
        stats_scrollbar.grid(row=0, column=1, sticky='ns')
        stats_frame.grid_rowconfigure(0, weight=1)
        stats_frame.grid_columnconfigure(0, weight=1)
        
        # 更新统计数据
        self.update_settings_statistics()

    def change_language(self, event=None):
        """更改语言"""
        new_lang = self.language_var.get()
        if new_lang != self.language:
            self.switch_language(new_lang)
            # 重启应用以应用新语言
            messagebox.showinfo(self.get_text('success'), '语言已更改，请重启应用以完全生效更改。')
    
    def update_settings_statistics(self):
        """更新设置标签页的统计数据"""
        # 清空文本框
        self.settings_stats_text.delete(1.0, tk.END)
        
        # 统计数据
        total_count = len(self.data)
        subject_counts = {}
        difficulty_counts = {}
        total_size = os.path.getsize(self.data_file) if os.path.exists(self.data_file) else 0
        
        for item in self.data:
            # 统计科目
            subject = item['科目']
            subject_counts[subject] = subject_counts.get(subject, 0) + 1
            
            # 统计难度
            difficulty = item['难度']
            difficulty_counts[difficulty] = difficulty_counts.get(difficulty, 0) + 1
        
        # 生成统计文本
        stats_text = f"📊 数据统计\n"
        stats_text += f"总错题数: {total_count}\n"
        stats_text += f"数据文件大小: {total_size} 字节\n"
        stats_text += f"上次更新: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        
        stats_text += "各科目错题数:\n"
        for subject, count in sorted(subject_counts.items()):
            stats_text += f"  {subject}: {count}题\n"
        
        stats_text += "\n各难度错题数:\n"
        for difficulty, count in sorted(difficulty_counts.items()):
            stats_text += f"  {difficulty}: {count}题\n"
        
        # 显示统计信息
        self.settings_stats_text.insert(tk.END, stats_text)
        self.settings_stats_text.config(state=tk.DISABLED)  # 设置为只读

    def clear_all_data(self):
        """清空所有数据"""
        if messagebox.askyesno("确认", "确定要清空所有错题数据吗？此操作不可恢复！"):
            self.data = []
            self.save_data()
            self.refresh_data()
            self.update_search_results()
            self.update_statistics()
            self.update_settings_statistics()
            messagebox.showinfo("成功", "所有数据已清空！")

    def backup_data(self):
        """备份数据"""
        # 询问备份位置
        backup_dir = filedialog.askdirectory(title="选择备份位置")
        if not backup_dir:
            return
        
        try:
            import shutil
            backup_file = os.path.join(backup_dir, f"mistakebook_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv")
            shutil.copy2(self.data_file, backup_file)
            messagebox.showinfo("成功", f"数据已备份到: {backup_file}")
        except Exception as e:
            messagebox.showerror("错误", f"备份失败: {str(e)}")

    def restore_data(self):
        """恢复数据"""
        # 询问备份文件
        backup_file = filedialog.askopenfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
            title="选择备份文件进行恢复"
        )
        
        if not backup_file:
            return
        
        try:
            # 确认恢复操作
            if not messagebox.askyesno("确认", "确定要从备份文件恢复数据吗？当前数据将被覆盖！"):
                return
                
            import shutil
            shutil.copy2(backup_file, self.data_file)
            self.load_data()  # 重新加载数据
            self.refresh_data()
            self.update_search_results()
            self.update_statistics()
            self.update_settings_statistics()
            messagebox.showinfo("成功", "数据已从备份恢复！")
        except Exception as e:
            messagebox.showerror("错误", f"恢复失败: {str(e)}")



    def show_about(self):
        """显示关于信息"""
        about_text = """错题本管理系统 v1.0

作者: Zhang Yangyi
功能: 管理错题，支持导入导出多种格式
界面: 使用tkinter构建

感谢使用本系统！"""
        messagebox.showinfo("关于", about_text)

    def show_help(self):
        """显示帮助信息"""
        # 创建帮助窗口
        help_window = tk.Toplevel(self.root)
        help_window.title("帮助")
        help_window.geometry("800x600")
        help_window.transient(self.root)
        help_window.grab_set()  # 模态窗口
        
        # 创建主框架
        main_frame = ttk.Frame(help_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        title_label = tk.Label(main_frame, text="错题本管理系统使用帮助", 
                              font=('Microsoft YaHei', 16, 'bold'))
        title_label.pack(anchor=tk.W, pady=(0, 20))
        
        # 创建文本框和滚动条
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        help_text_widget = tk.Text(text_frame, wrap=tk.WORD, 
                                  font=('Microsoft YaHei', 12),  # 增大字体
                                  padx=10, pady=10, spacing1=5, spacing3=5)
        scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=help_text_widget.yview)
        help_text_widget.configure(yscrollcommand=scrollbar.set)
        
        # 插入详细的帮助内容
        detailed_help_content = """错题本管理系统使用帮助

一、系统概述
错题本管理系统是一个功能全面的学习辅助工具，帮助用户高效管理错题，提供错题录入、分类、检索、分析等功能。

二、功能模块详解

1. 错题列表标签页
   - 功能：查看、管理所有错题记录
   - 添加错题：点击"添加"按钮，在弹窗中填写题目信息
     * 时间：自动填入当前时间，也可手动修改
     * 科目：选择题目所属科目（语文、数学、英语等）
     * 题干：输入题目内容
     * 正确答案：填写正确解题思路或答案
     * 附件：可关联相关文件（图片、文档等）
     * 难度：标记题目难度（简单、中等、困难）
   - 编辑错题：选中错题，点击"编辑"按钮进行修改
   - 删除错题：选中错题，点击"删除"按钮
   - 查看详情：点击"详情"按钮查看完整题目信息

2. 搜索标签页
   - 多条件筛选：支持按科目、难度、关键词、日期范围筛选
   - 智能搜索：输入关键词可同时搜索题干、答案、科目等字段
   - 操作功能：对搜索结果可执行查看、编辑、删除等操作
   - 组合筛选：可同时应用多个筛选条件获得精确结果

3. 导入导出标签页
   - 导入功能：
     * CSV格式：支持标准CSV格式数据导入
     * TXT格式：支持特定格式文本导入
     * Excel格式：支持XLSX格式导入（需安装pandas）
     * Word格式：支持DOCX格式导入（需安装python-docx）
   - 导出功能：
     * CSV格式：导出为标准CSV文件，便于Excel打开
     * TXT格式：导出为文本格式，便于阅读
     * Excel格式：导出为XLSX格式（需安装pandas）
     * Word格式：导出为DOCX格式（需安装python-docx）
     * PDF格式：导出为PDF文档（需安装reportlab）
   - 批量操作：支持导出全部数据或仅选中数据

4. 合集管理标签页（新增）
   - 创建合集：可创建具有名称和描述的错题合集
   - 管理合集：查看、编辑、删除已创建的合集
   - 题目分配：可将错题分配到特定合集
   - 导出合集：可将指定合集单独导出为文件

5. 统计标签页
   - 数据概览：显示总体错题数、各科目分布、各难度分布
   - 可视化图表：以饼图和柱状图形式展示数据分布
   - 报告生成：可导出统计报告和图表
   - 实时更新：数据变化时自动更新统计信息

6. 设置标签页
   - 主题切换：支持浅色/深色主题切换
   - 字体设置：可调整界面字体大小
   - 数据管理：提供数据备份与恢复功能
   - 应用信息：查看版本和关于信息

三、使用技巧
   1. 建议定期备份数据，防止数据丢失
   2. 合理使用难度标记，便于后续复习
   3. 利用搜索功能快速定位特定错题
   4. 通过统计功能分析学习薄弱环节
   5. 利用合集功能将相关错题归类整理

四、快捷操作
   - 选中错题后可直接编辑或查看详情
   - 可同时选中多条错题进行批量操作
   - 搜索功能支持实时筛选，输入即响应
   - 支持拖拽调整窗口大小以获得更好的显示效果

五、常见问题
   Q: 无法导入Excel文件？
   A: 请确保已安装pandas库：pip install pandas openpyxl

   Q: 导出PDF失败？
   A: 请确保已安装reportlab库：pip install reportlab

   Q: 如何备份数据？
   A: 在设置标签页中点击"数据备份"按钮，选择备份位置即可

   Q: 可以将错题按主题分类吗？
   A: 可以使用合集功能创建主题合集，将相关错题归类管理
"""
        help_text_widget.insert(tk.END, detailed_help_content)
        help_text_widget.config(state=tk.DISABLED)  # 设置为只读
        
        # 布局
        help_text_widget.grid(row=0, column=0, sticky='nsew')
        scrollbar.grid(row=0, column=1, sticky='ns')
        
        text_frame.grid_rowconfigure(0, weight=1)
        text_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_rowconfigure(1, weight=1)
        main_frame.grid_columnconfigure(0, weight=1)
        
        # 添加关闭按钮
        close_button = ttk.Button(main_frame, text="关闭", 
                                 command=help_window.destroy)
        close_button.pack(pady=(10, 0))

    def update_statistics(self):
        """更新统计数据"""
        # 清空文本框
        self.stats_text.delete(1.0, tk.END)
        
        # 统计数据
        total_count = len(self.data)
        subject_counts = {}
        difficulty_counts = {}
        
        for item in self.data:
            # 统计科目
            subject = item['科目']
            subject_counts[subject] = subject_counts.get(subject, 0) + 1
            
            # 统计难度
            difficulty = item['难度']
            difficulty_counts[difficulty] = difficulty_counts.get(difficulty, 0) + 1
        
        # 生成统计文本
        stats_text = f"📊 数据统计\n"
        stats_text += f"总错题数: {total_count}\n\n"
        
        stats_text += "各科目错题数:\n"
        for subject, count in subject_counts.items():
            stats_text += f"  {subject}: {count}题\n"
        
        stats_text += "\n各难度错题数:\n"
        for difficulty, count in difficulty_counts.items():
            stats_text += f"  {difficulty}: {count}题\n"
        
        # 显示统计信息
        self.stats_text.insert(tk.END, stats_text)
        self.stats_text.config(state=tk.DISABLED)  # 设置为只读

    def export_to_txt(self):
        """导出数据到TXT"""
        # 询问保存位置
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            title="导出为TXT"
        )
        
        if not file_path:
            return
        
        # 写入数据到TXT文件，使用UTF-8编码
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("错题本导出\n")
            f.write("="*50 + "\n\n")
            
            for i, row in enumerate(self.data, 1):
                f.write(f"第{i}题\n")
                f.write(f"时间: {row['时间']}\n")
                f.write(f"科目: {row['科目']}\n")
                f.write(f"难度: {row['难度']}\n")
                f.write(f"题干: {row['题干']}\n")
                f.write(f"正确答案: {row['正确答案']}\n")
                f.write(f"附件路径: {row['附件路径'] if row['附件路径'] else '无'}\n")
                f.write("-" * 50 + "\n\n")
        
        messagebox.showinfo("成功", f"数据已导出到 {file_path}")

    def import_from_txt(self):
        """从TXT文件导入数据"""
        # 询问选择TXT文件
        file_path = filedialog.askopenfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            title="从TXT导入"
        )
        
        if not file_path:
            return
        
        try:
            # 读取TXT文件内容
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 解析TXT内容并导入数据
            imported_data = self.parse_txt_content(content)
            
            if imported_data:
                # 添加到现有数据中
                self.data.extend(imported_data)
                
                # 保存到CSV文件
                self.save_data()
                
                # 刷新界面
                self.refresh_data()
                
                messagebox.showinfo("成功", f"成功导入 {len(imported_data)} 条错题！")
            else:
                messagebox.showwarning("警告", "未找到有效的错题数据，请检查TXT文件格式。")
                
        except Exception as e:
            messagebox.showerror("错误", f"导入失败: {str(e)}")

    def parse_txt_content(self, content):
        """解析TXT文件内容，返回错题数据列表"""
        lines = content.split('\n')
        imported_data = []
        
        # 用于存储当前错题信息的字典
        current_item = {}
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # 检查是否是新错题的开始
            if line.startswith("第") and "题" in line:
                # 如果已有current_item，保存之前的数据
                if current_item and '时间' in current_item:
                    # 设置默认值
                    if '附件路径' not in current_item:
                        current_item['附件路径'] = ''
                    if '难度' not in current_item:
                        current_item['难度'] = '中等'
                    imported_data.append(current_item)
                
                # 开始新的错题记录
                current_item = {}
                
            elif line.startswith("时间:"):
                current_item['时间'] = line[3:].strip()  # 去除"时间:"前缀
                
            elif line.startswith("科目:"):
                current_item['科目'] = line[3:].strip()
                
            elif line.startswith("难度:"):
                current_item['难度'] = line[3:].strip()
                
            elif line.startswith("题干:"):
                # 题干可能有多行，需要处理
                current_item['题干'] = line[3:].strip()
                i += 1
                # 继续读取直到遇到下一个字段或分隔线
                while i < len(lines):
                    next_line = lines[i].strip()
                    if (next_line.startswith("正确答案:") or 
                        next_line.startswith("附件路径:") or
                        next_line.startswith("第") and "题" in next_line or
                        next_line.startswith("-" * 10)):  # 分隔线
                        i -= 1  # 回退一行，因为当前行不属于题干
                        break
                    else:
                        # 添加到题干中
                        if '题干' in current_item:
                            current_item['题干'] += "\n" + next_line
                        else:
                            current_item['题干'] = next_line
                    i += 1
                    
            elif line.startswith("正确答案:"):
                # 正确答案可能有多行，需要处理
                current_item['正确答案'] = line[5:].strip()
                i += 1
                # 继续读取直到遇到下一个字段或分隔线
                while i < len(lines):
                    next_line = lines[i].strip()
                    if (next_line.startswith("附件路径:") or
                        next_line.startswith("第") and "题" in next_line or
                        next_line.startswith("-" * 10)):  # 分隔线
                        i -= 1  # 回退一行，因为当前行不属于正确答案
                        break
                    else:
                        # 添加到正确答案中
                        if '正确答案' in current_item:
                            current_item['正确答案'] += "\n" + next_line
                        else:
                            current_item['正确答案'] = next_line
                    i += 1
                    
            elif line.startswith("附件路径:"):
                current_item['附件路径'] = line[5:].strip()
            
            i += 1
        
        # 添加最后一个错题
        if current_item and '时间' in current_item:
            # 设置默认值
            if '附件路径' not in current_item:
                current_item['附件路径'] = ''
            if '难度' not in current_item:
                current_item['难度'] = '中等'
                
            imported_data.append(current_item)
        
        return imported_data

    def create_statistics_tab(self):
        """创建统计标签页"""
        # 创建框架
        self.statistics_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.statistics_frame, text="📊 统计")
        
        # 统计信息显示区域
        info_frame = ttk.LabelFrame(self.statistics_frame, text="📈 统计信息", padding=10)
        info_frame.pack(fill=tk.X, pady=5)
        
        # 统计信息文本框
        self.stats_info_text = tk.Text(info_frame, height=8, wrap=tk.WORD, font=('Microsoft YaHei', 10))
        stats_info_scrollbar = ttk.Scrollbar(info_frame, orient=tk.VERTICAL, command=self.stats_info_text.yview)
        self.stats_info_text.configure(yscrollcommand=stats_info_scrollbar.set)
        
        self.stats_info_text.grid(row=0, column=0, sticky='nsew')
        stats_info_scrollbar.grid(row=0, column=1, sticky='ns')
        info_frame.grid_rowconfigure(0, weight=1)
        info_frame.grid_columnconfigure(0, weight=1)
        
        # 图表显示区域
        chart_frame = ttk.LabelFrame(self.statistics_frame, text="📊 统计图表", padding=10)
        chart_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 创建画布用于显示图表
        self.chart_canvas = tk.Canvas(chart_frame, bg='white')
        chart_canvas_scrollbar = ttk.Scrollbar(chart_frame, orient=tk.VERTICAL, command=self.chart_canvas.yview)
        self.chart_canvas.configure(yscrollcommand=chart_canvas_scrollbar.set)
        
        # 创建第二个滚动条用于水平滚动
        chart_h_scrollbar = ttk.Scrollbar(chart_frame, orient=tk.HORIZONTAL, command=self.chart_canvas.xview)
        self.chart_canvas.configure(xscrollcommand=chart_h_scrollbar.set)
        
        # 将画布和滚动条布局
        self.chart_canvas.grid(row=0, column=0, sticky='nsew')
        chart_canvas_scrollbar.grid(row=0, column=1, sticky='ns')
        chart_h_scrollbar.grid(row=1, column=0, sticky='ew')
        
        chart_frame.grid_rowconfigure(0, weight=1)
        chart_frame.grid_columnconfigure(0, weight=1)
        
        # 按钮区域
        button_frame = ttk.Frame(self.statistics_frame)
        button_frame.pack(fill=tk.X, pady=5)
        
        # 创建样式
        style = ttk.Style()
        style.configure('Action.TButton', font=('Microsoft YaHei', 10))
        
        # 刷新统计信息按钮
        ttk.Button(button_frame, text="🔄 刷新统计", command=self.refresh_statistics, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=2)
        
        # 导出统计图表按钮
        ttk.Button(button_frame, text="📤 "+self.get_text('export_charts'), command=self.export_statistics_chart, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=2)
        
        # 导出统计报告按钮(TXT)
        ttk.Button(button_frame, text="📝 "+self.get_text('export_report_txt'), command=self.export_statistics_report, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=2)
        
        # 导出统计报告按钮(PDF)
        ttk.Button(button_frame, text="📄 "+self.get_text('export_report_pdf'), command=self.export_statistics_report_pdf, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=2)
        
        # 导出统计图表按钮(PDF)
        ttk.Button(button_frame, text="📊 "+self.get_text('menu_export_stats_charts_pdf'), command=self.export_statistics_charts_pdf, 
                  style='Action.TButton').pack(side=tk.LEFT, padx=2)
        
        # 配置列权重，使按钮平均分布
        for i in range(5):  # 更新为5个按钮
            button_frame.columnconfigure(i, weight=1)
        
        # 初始化统计信息
        self.refresh_statistics()

    def refresh_statistics(self):
        """刷新统计信息和图表"""
        # 更新统计信息文本
        self.update_statistics_info()
        
        # 更新图表
        self.update_statistics_chart()

    def update_statistics_info(self):
        """更新统计信息文本"""
        # 清空文本框
        self.stats_info_text.delete(1.0, tk.END)
        
        # 统计数据
        total_count = len(self.data)
        subject_counts = {}
        difficulty_counts = {}
        
        for item in self.data:
            # 统计科目
            subject = item['科目']
            subject_counts[subject] = subject_counts.get(subject, 0) + 1
            
            # 统计难度
            difficulty = item['难度']
            difficulty_counts[difficulty] = difficulty_counts.get(difficulty, 0) + 1
        
        # 生成统计文本
        stats_text = f"📊 错题统计报告\n"
        stats_text += f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        stats_text += f"总错题数: {total_count}\n\n"
        
        stats_text += "各科目错题数:\n"
        for subject, count in sorted(subject_counts.items(), key=lambda x: x[1], reverse=True):
            percentage = (count / total_count * 100) if total_count > 0 else 0
            stats_text += f"  {subject}: {count}题 ({percentage:.1f}%)\n"
        
        stats_text += "\n各难度错题数:\n"
        for difficulty, count in sorted(difficulty_counts.items(), key=lambda x: x[1], reverse=True):
            percentage = (count / total_count * 100) if total_count > 0 else 0
            stats_text += f"  {difficulty}: {count}题 ({percentage:.1f}%)\n"
        
        # 显示统计信息
        self.stats_info_text.insert(tk.END, stats_text)
        self.stats_info_text.config(state=tk.DISABLED)  # 设置为只读

    def update_statistics_chart(self):
        """更新统计图表"""
        # 清空画布
        self.chart_canvas.delete("all")
        
        if not self.data:
            self.chart_canvas.create_text(400, 100, text="暂无数据", font=('Microsoft YaHei', 14))
            return
        
        # 统计科目和难度数据
        subject_counts = {}
        difficulty_counts = {}
        
        for item in self.data:
            # 统计科目
            subject = item['科目']
            subject_counts[subject] = subject_counts.get(subject, 0) + 1
            
            # 统计难度
            difficulty = item['难度']
            difficulty_counts[difficulty] = difficulty_counts.get(difficulty, 0) + 1
        
        # 画布尺寸
        canvas_width = 1600
        canvas_height = 1400  # 增加高度以容纳所有图表
        
        # 设置画布滚动区域
        self.chart_canvas.config(scrollregion=(0, 0, canvas_width, canvas_height))
        
        # 绘制标题
        self.chart_canvas.create_text(800, 30, text="错题统计图表", font=('Microsoft YaHei', 18, 'bold'))
        
        # 绘制科目分布饼图
        self.draw_subject_pie_chart(subject_counts, 50, 80, 500, 500)
        
        # 绘制难度分布饼图
        self.draw_difficulty_pie_chart(difficulty_counts, 700, 80, 500, 500)
        
        # 绘制科目分布柱状图
        self.draw_subject_bar_chart(subject_counts, 50, 650, 700, 500)
        
        # 绘制难度分布柱状图
        self.draw_difficulty_bar_chart(difficulty_counts, 850, 650, 500, 500)

    def draw_subject_pie_chart(self, subject_counts, x, y, width, height):
        """绘制科目分布饼图"""
        # 绘制标题
        self.chart_canvas.create_text(x + width//2, y - 10, text="科目分布饼图", font=('Microsoft YaHei', 12, 'bold'))
        
        if not subject_counts:
            self.chart_canvas.create_text(x + width//2, y + height//2, text="暂无数据", font=('Microsoft YaHei', 10))
            return
        
        # 颜色列表
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#DDA0DD', '#98D8C8', '#F7DC6F', '#BB8FCE']
        
        # 计算总数
        total = sum(subject_counts.values())
        
        # 绘制饼图
        start_angle = 0
        center_x = x + width // 2
        center_y = y + height // 2
        radius = min(width, height) // 2 - 20
        
        i = 0
        for subject, count in subject_counts.items():
            angle = 360 * count / total if total > 0 else 0
            color = colors[i % len(colors)]
            
            # 绘制扇形
            self.chart_canvas.create_arc(
                center_x - radius, center_y - radius,
                center_x + radius, center_y + radius,
                start=start_angle, extent=angle,
                fill=color, outline='white', width=2
            )
            
            # 计算标签位置
            label_angle = start_angle + angle / 2
            import math
            label_x = center_x + radius * 0.7 * math.cos(math.radians(label_angle))
            label_y = center_y + radius * 0.7 * math.sin(math.radians(label_angle))
            
            # 绘制标签
            self.chart_canvas.create_text(label_x, label_y, text=f"{subject}\n{count}题", 
                                        font=('Microsoft YaHei', 8), fill='black')
            
            start_angle += angle
            i += 1
        
        # 绘制图例
        legend_x = x + width + 20
        legend_y = y
        i = 0
        for subject, count in subject_counts.items():
            color = colors[i % len(colors)]
            self.chart_canvas.create_rectangle(legend_x, legend_y, legend_x + 15, legend_y + 15, 
                                             fill=color, outline='black')
            percentage = (count / total * 100) if total > 0 else 0
            self.chart_canvas.create_text(legend_x + 20, legend_y + 8, text=f"{subject}: {count}题 ({percentage:.1f}%)", 
                                        font=('Microsoft YaHei', 8), anchor=tk.W)
            legend_y += 20
            i += 1

    def draw_difficulty_pie_chart(self, difficulty_counts, x, y, width, height):
        """绘制难度分布饼图"""
        # 绘制标题
        self.chart_canvas.create_text(x + width//2, y - 10, text="难度分布饼图", font=('Microsoft YaHei', 12, 'bold'))
        
        if not difficulty_counts:
            self.chart_canvas.create_text(x + width//2, y + height//2, text="暂无数据", font=('Microsoft YaHei', 10))
            return
        
        # 颜色列表
        colors = ['#FF9F43', '#10AC84', '#EE5A24']
        
        # 计算总数
        total = sum(difficulty_counts.values())
        
        # 绘制饼图
        start_angle = 0
        center_x = x + width // 2
        center_y = y + height // 2
        radius = min(width, height) // 2 - 20
        
        i = 0
        for difficulty, count in difficulty_counts.items():
            angle = 360 * count / total if total > 0 else 0
            color = colors[i % len(colors)]
            
            # 绘制扇形
            self.chart_canvas.create_arc(
                center_x - radius, center_y - radius,
                center_x + radius, center_y + radius,
                start=start_angle, extent=angle,
                fill=color, outline='white', width=2
            )
            
            # 计算标签位置
            label_angle = start_angle + angle / 2
            import math
            label_x = center_x + radius * 0.7 * math.cos(math.radians(label_angle))
            label_y = center_y + radius * 0.7 * math.sin(math.radians(label_angle))
            
            # 绘制标签
            self.chart_canvas.create_text(label_x, label_y, text=f"{difficulty}\n{count}题", 
                                        font=('Microsoft YaHei', 8), fill='black')
            
            start_angle += angle
            i += 1
        
        # 绘制图例
        legend_x = x + width + 20
        legend_y = y
        i = 0
        for difficulty, count in difficulty_counts.items():
            color = colors[i % len(colors)]
            self.chart_canvas.create_rectangle(legend_x, legend_y, legend_x + 15, legend_y + 15, 
                                             fill=color, outline='black')
            percentage = (count / total * 100) if total > 0 else 0
            self.chart_canvas.create_text(legend_x + 20, legend_y + 8, text=f"{difficulty}: {count}题 ({percentage:.1f}%)", 
                                        font=('Microsoft YaHei', 8), anchor=tk.W)
            legend_y += 20
            i += 1

    def draw_subject_bar_chart(self, subject_counts, x, y, width, height):
        """绘制科目分布柱状图"""
        # 绘制标题
        self.chart_canvas.create_text(x + width//2, y - 10, text="科目分布柱状图", font=('Microsoft YaHei', 12, 'bold'))
        
        if not subject_counts:
            self.chart_canvas.create_text(x + width//2, y + height//2, text="暂无数据", font=('Microsoft YaHei', 10))
            return
        
        # 颜色列表
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#DDA0DD', '#98D8C8', '#F7DC6F', '#BB8FCE']
        
        # 获取排序后的科目和数量
        sorted_subjects = sorted(subject_counts.items(), key=lambda x: x[1], reverse=True)
        subjects = [item[0] for item in sorted_subjects]
        counts = [item[1] for item in sorted_subjects]
        
        if not counts:
            return
        
        max_count = max(counts) if counts else 1
        
        # 图表边距
        margin = 50
        chart_width = width - 2 * margin
        chart_height = height - 2 * margin
        
        # 计算柱子的宽度和间距
        num_bars = len(subjects)
        if num_bars > 0:
            bar_width = min(40, chart_width // num_bars * 0.8)  # 最大宽度40像素
            spacing = (chart_width - num_bars * bar_width) // (num_bars + 1) if num_bars > 0 else 0
            
            # 绘制坐标轴
            self.chart_canvas.create_line(x + margin, y + height - margin, 
                                        x + width - margin, y + height - margin, width=2)  # X轴
            self.chart_canvas.create_line(x + margin, y + margin, 
                                        x + margin, y + height - margin, width=2)  # Y轴
            
            # 绘制刻度和标签
            # Y轴刻度
            for i in range(0, max_count + 1, max(1, max_count // 5)):
                y_pos = y + height - margin - (i / max_count) * chart_height
                self.chart_canvas.create_line(x + margin - 5, y_pos, x + margin, y_pos, width=1)
                self.chart_canvas.create_text(x + margin - 10, y_pos, text=str(i), font=('Microsoft YaHei', 8), anchor=tk.E)
            
            # 绘制柱子和标签
            for i, (subject, count) in enumerate(sorted_subjects):
                color = colors[i % len(colors)]
                
                # 计算柱子位置
                bar_x = x + margin + spacing + i * (bar_width + spacing)
                bar_height = (count / max_count) * chart_height if max_count > 0 else 0
                bar_y = y + height - margin - bar_height
                
                # 绘制柱子
                self.chart_canvas.create_rectangle(bar_x, bar_y, bar_x + bar_width, y + height - margin,
                                                fill=color, outline='black', width=1)
                
                # 绘制数值标签
                self.chart_canvas.create_text(bar_x + bar_width/2, bar_y - 5, text=str(count), 
                                            font=('Microsoft YaHei', 8), anchor=tk.S)
                
                # 绘制科目标签
                self.chart_canvas.create_text(bar_x + bar_width/2, y + height - margin + 15, text=subject, 
                                            font=('Microsoft YaHei', 8), anchor=tk.N)

    def draw_difficulty_bar_chart(self, difficulty_counts, x, y, width, height):
        """绘制难度分布柱状图"""
        # 绘制标题
        self.chart_canvas.create_text(x + width//2, y - 10, text="难度分布柱状图", font=('Microsoft YaHei', 12, 'bold'))
        
        if not difficulty_counts:
            self.chart_canvas.create_text(x + width//2, y + height//2, text="暂无数据", font=('Microsoft YaHei', 10))
            return
        
        # 颜色列表
        colors = ['#FF9F43', '#10AC84', '#EE5A24']
        
        # 获取排序后的难度和数量
        sorted_difficulties = sorted(difficulty_counts.items(), key=lambda x: x[1], reverse=True)
        difficulties = [item[0] for item in sorted_difficulties]
        counts = [item[1] for item in sorted_difficulties]
        
        if not counts:
            return
        
        max_count = max(counts) if counts else 1
        
        # 图表边距
        margin = 50
        chart_width = width - 2 * margin
        chart_height = height - 2 * margin
        
        # 计算柱子的宽度和间距
        num_bars = len(difficulties)
        if num_bars > 0:
            bar_width = min(60, chart_width // num_bars * 0.8)  # 最大宽度60像素
            spacing = (chart_width - num_bars * bar_width) // (num_bars + 1) if num_bars > 0 else 0
            
            # 绘制坐标轴
            self.chart_canvas.create_line(x + margin, y + height - margin, 
                                        x + width - margin, y + height - margin, width=2)  # X轴
            self.chart_canvas.create_line(x + margin, y + margin, 
                                        x + margin, y + height - margin, width=2)  # Y轴
            
            # 绘制刻度和标签
            # Y轴刻度
            for i in range(0, max_count + 1, max(1, max_count // 5)):
                y_pos = y + height - margin - (i / max_count) * chart_height
                self.chart_canvas.create_line(x + margin - 5, y_pos, x + margin, y_pos, width=1)
                self.chart_canvas.create_text(x + margin - 10, y_pos, text=str(i), font=('Microsoft YaHei', 8), anchor=tk.E)
            
            # 绘制柱子和标签
            for i, (difficulty, count) in enumerate(sorted_difficulties):
                color = colors[i % len(colors)]
                
                # 计算柱子位置
                bar_x = x + margin + spacing + i * (bar_width + spacing)
                bar_height = (count / max_count) * chart_height if max_count > 0 else 0
                bar_y = y + height - margin - bar_height
                
                # 绘制柱子
                self.chart_canvas.create_rectangle(bar_x, bar_y, bar_x + bar_width, y + height - margin,
                                                fill=color, outline='black', width=1)
                
                # 绘制数值标签
                self.chart_canvas.create_text(bar_x + bar_width/2, bar_y - 5, text=str(count), 
                                            font=('Microsoft YaHei', 8), anchor=tk.S)
                
                # 绘制难度标签
                self.chart_canvas.create_text(bar_x + bar_width/2, y + height - margin + 15, text=difficulty, 
                                            font=('Microsoft YaHei', 8), anchor=tk.N)

    def export_statistics_chart(self):
        """导出统计图表"""
        try:
            from PIL import Image, ImageDraw
            import math
            
            # 询问保存位置
            file_path = filedialog.asksaveasfilename(
                defaultextension=".png",
                filetypes=[("PNG files", "*.png"), ("JPEG files", "*.jpg"), ("All files", "*.*")],
                title="导出统计图表"
            )
            
            if not file_path:
                return
            
            # 统计数据
            subject_counts = {}
            difficulty_counts = {}
            
            for item in self.data:
                # 统计科目
                subject = item['科目']
                subject_counts[subject] = subject_counts.get(subject, 0) + 1
                
                # 统计难度
                difficulty = item['难度']
                difficulty_counts[difficulty] = difficulty_counts.get(difficulty, 0) + 1
            
            # 创建图像
            img_width, img_height = 1200, 800
            img = Image.new('RGB', (img_width, img_height), 'white')
            draw = ImageDraw.Draw(img)
            
            # 绘制标题
            try:
                from PIL import ImageFont
                font_large = ImageFont.truetype("arial.ttf", 24)
                font_medium = ImageFont.truetype("arial.ttf", 16)
                font_small = ImageFont.truetype("arial.ttf", 12)
            except:
                # 如果找不到字体文件，则使用默认字体
                font_large = None
                font_medium = None
                font_small = None
            
            # 添加标题
            title = "错题统计图表"
            draw.text((img_width//2 - 80, 20), title, fill='black', font=font_large)
            
            # 绘制科目分布饼图
            self.draw_pil_pie_chart(draw, subject_counts, 100, 80, 300, 300, "科目分布", font_small)
            
            # 绘制难度分布饼图
            self.draw_pil_pie_chart(draw, difficulty_counts, 450, 80, 300, 300, "难度分布", font_small)
            
            # 绘制科目分布柱状图
            self.draw_pil_bar_chart(draw, subject_counts, 100, 420, 500, 300, "科目分布柱状图", font_small)
            
            # 绘制难度分布柱状图
            self.draw_pil_bar_chart(draw, difficulty_counts, 650, 420, 400, 300, "难度分布柱状图", font_small)
            
            # 保存图像
            img.save(file_path)
            messagebox.showinfo("成功", f"统计图表已导出到 {file_path}")
            
        except ImportError:
            messagebox.showerror("错误", "需要安装Pillow库，请运行: pip install Pillow")
        except Exception as e:
            messagebox.showerror("错误", f"导出图表失败: {str(e)}")

    def draw_pil_pie_chart(self, draw, data, x, y, width, height, title, font):
        """使用Pillow绘制饼图"""
        if not data:
            draw.text((x + width//2 - 30, y + height//2), "暂无数据", fill='black', font=font)
            return
        
        # 颜色列表
        colors = [(255, 107, 107), (78, 205, 196), (69, 183, 209), (150, 206, 180), 
                  (255, 234, 167), (221, 160, 221), (152, 216, 200), (247, 220, 111), (187, 143, 206)]
        
        # 计算总数
        total = sum(data.values())
        
        # 添加标题
        draw.text((x + width//2 - len(title)*5, y - 25), title, fill='black', font=font)
        
        # 绘制饼图
        start_angle = 0
        center_x = x + width // 2
        center_y = y + height // 2
        radius = min(width, height) // 2 - 20
        
        i = 0
        for item, count in data.items():
            angle = 360 * count / total if total > 0 else 0
            color = colors[i % len(colors)]
            
            # 将角度转换为PIL格式 (start, end) 以度为单位, 0度在3点钟位置, 顺时针为正
            start_deg = start_angle - 90  # 转换为PIL标准
            end_deg = start_deg + angle
            
            # 绘制扇形
            bbox = [center_x - radius, center_y - radius, center_x + radius, center_y + radius]
            draw.pieslice(bbox, start_deg, end_deg, fill=color, outline='white')
            
            # 计算标签位置
            label_angle = start_angle + angle / 2
            import math
            label_x = center_x + radius * 0.7 * math.cos(math.radians(label_angle))
            label_y = center_y + radius * 0.7 * math.sin(math.radians(label_angle))
            
            # 绘制标签
            draw.text((label_x - 10, label_y - 5), f"{item}\n{count}", fill='black', font=font)
            
            start_angle += angle
            i += 1

    def draw_pil_bar_chart(self, draw, data, x, y, width, height, title, font):
        """使用Pillow绘制柱状图"""
        if not data:
            draw.text((x + width//2 - 30, y + height//2), "暂无数据", fill='black', font=font)
            return
        
        # 颜色列表
        colors = [(255, 107, 107), (78, 205, 196), (69, 183, 209), (150, 206, 180), 
                  (255, 234, 167), (221, 160, 221), (152, 216, 200), (247, 220, 111), (187, 143, 206)]
        
        # 获取排序后的数据
        sorted_data = sorted(data.items(), key=lambda x: x[1], reverse=True)
        items = [item[0] for item in sorted_data]
        counts = [item[1] for item in sorted_data]
        
        max_count = max(counts) if counts else 1
        
        # 图表边距
        margin = 50
        chart_width = width - 2 * margin
        chart_height = height - 2 * margin
        
        # 添加标题
        draw.text((x + width//2 - len(title)*5, y - 25), title, fill='black', font=font)
        
        # 绘制坐标轴
        draw.line([(x + margin, y + height - margin), (x + width - margin, y + height - margin)], fill='black', width=2)  # X轴
        draw.line([(x + margin, y + margin), (x + margin, y + height - margin)], fill='black', width=2)  # Y轴
        
        # 绘制刻度和标签
        # Y轴刻度
        for i in range(0, max_count + 1, max(1, max_count // 5)):
            y_pos = y + height - margin - (i / max_count) * chart_height
            draw.line([(x + margin - 5, y_pos), (x + margin, y_pos)], fill='black', width=1)
            if font:
                draw.text((x + margin - 20, y_pos - 5), str(i), fill='black', font=font)
        
        # 计算柱子的宽度和间距
        num_bars = len(items)
        if num_bars > 0:
            bar_width = min(40, chart_width // num_bars * 0.8)
            spacing = (chart_width - num_bars * bar_width) // (num_bars + 1) if num_bars > 0 else 0
            
            # 绘制柱子和标签
            for i, (item, count) in enumerate(sorted_data):
                color = colors[i % len(colors)]
                
                # 计算柱子位置
                bar_x = x + margin + spacing + i * (bar_width + spacing)
                bar_height = (count / max_count) * chart_height if max_count > 0 else 0
                bar_y = y + height - margin - bar_height
                
                # 绘制柱子
                draw.rectangle([bar_x, bar_y, bar_x + bar_width, y + height - margin], fill=color, outline='black')
                
                # 绘制数值标签
                if font:
                    draw.text((bar_x + bar_width/2 - 5, bar_y - 15), str(count), fill='black', font=font)
                
                # 绘制项目标签
                if font:
                    draw.text((bar_x + bar_width/2 - len(item)*3, y + height - margin + 5), item, fill='black', font=font)

    def export_statistics_charts_pdf(self):
        """导出统计图表为PDF"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import tempfile
            import os
            from PIL import Image as PILImage, ImageDraw
            import math
            
            # 注册中文字体
            font_name = 'SimSun'
            try:
                # 尝试注册字体
                pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))
            except:
                # 如果SimSun不可用，尝试其他常见中文字体
                try:
                    # 在Windows系统上查找常见字体文件
                    font_paths = [
                        r'C:\Windows\Fonts\simsun.ttc',    # 宋体
                        r'C:\Windows\Fonts\msyh.ttc',      # 微软雅黑
                        r'C:\Windows\Fonts\msyhbd.ttc',    # 微软雅黑粗体
                        r'C:\Windows\Fonts\simhei.ttf',    # 黑体
                    ]
                    
                    font_found = False
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            if 'simsun' in font_path.lower():
                                font_name = 'SimSun'
                            elif 'msyh' in font_path.lower():
                                font_name = 'MicrosoftYaHei'
                            elif 'simhei' in font_path.lower():
                                font_name = 'SimHei'
                            
                            pdfmetrics.registerFont(TTFont(font_name, font_path))
                            font_found = True
                            break
                    
                    if not font_found:
                        # 如果找不到系统字体，则报错
                        messagebox.showerror("错误", "未找到支持中文的字体文件")
                        return
                except:
                    messagebox.showerror("错误", "无法注册中文字体，请确保系统中有中文字体文件")
                    return
            
            # 询问保存位置
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
                title="导出统计图表为PDF"
            )
            
            if not file_path:
                return
            
            # 统计数据
            subject_counts = {}
            difficulty_counts = {}
            
            for item in self.data:
                # 统计科目
                subject = item['科目']
                subject_counts[subject] = subject_counts.get(subject, 0) + 1
                
                # 统计难度
                difficulty = item['难度']
                difficulty_counts[difficulty] = difficulty_counts.get(difficulty, 0) + 1
            
            # 创建临时图像文件
            img_width, img_height = 1200, 1600
            img = PILImage.new('RGB', (img_width, img_height), 'white')
            draw = ImageDraw.Draw(img)
            
            # 尝试加载字体
            try:
                from PIL import ImageFont
                font_large = ImageFont.truetype("arial.ttf", 24)
                font_medium = ImageFont.truetype("arial.ttf", 16)
                font_small = ImageFont.truetype("arial.ttf", 12)
            except:
                # 如果找不到字体文件，则使用默认字体
                font_large = None
                font_medium = None
                font_small = None
            
            # 添加标题
            title = "📊 错题统计图表"
            if font_large:
                draw.text((img_width//2 - 80, 20), title, fill='black', font=font_large)
            else:
                draw.text((img_width//2 - 80, 20), title, fill='black')
            
            # 绘制科目分布饼图
            self.draw_pil_pie_chart(draw, subject_counts, 50, 80, 500, 500, "科目分布饼图", font_medium)
            
            # 绘制难度分布饼图
            self.draw_pil_pie_chart(draw, difficulty_counts, 650, 80, 500, 500, "难度分布饼图", font_medium)
            
            # 绘制科目分布柱状图
            self.draw_pil_bar_chart(draw, subject_counts, 50, 650, 700, 500, "科目分布柱状图", font_medium)
            
            # 绘制难度分布柱状图
            self.draw_pil_bar_chart(draw, difficulty_counts, 750, 650, 400, 500, "难度分布柱状图", font_medium)
            
            # 保存临时图像
            temp_img_path = os.path.join(tempfile.gettempdir(), "temp_chart.png")
            img.save(temp_img_path)
            
            # 创建PDF文档
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            story = []
            
            # 标题样式（使用中文字体）
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=getSampleStyleSheet()['Title'],
                fontName=font_name,
                fontSize=16,
                spaceAfter=30,
                alignment=1  # 居中
            )
            
            # 正文样式（使用中文字体）
            content_style = ParagraphStyle(
                'CustomContent',
                parent=getSampleStyleSheet()['Normal'],
                fontName=font_name,
                fontSize=10,
                spaceAfter=12
            )
            
            # 添加标题
            title_para = Paragraph("📊 错题统计图表", title_style)
            story.append(title_para)
            story.append(Spacer(1, 0.2 * inch))
            
            # 添加生成时间
            time_para = Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", content_style)
            story.append(time_para)
            story.append(Spacer(1, 0.2 * inch))
            
            # 添加图表
            try:
                # 缩放图像以适应PDF页面
                img_obj = Image(temp_img_path, width=6*inch, height=8*inch)
                img_obj.hAlign = 'CENTER'
                story.append(img_obj)
            except:
                # 如果无法添加图像，则添加说明文字
                story.append(Paragraph("图表无法显示", content_style))
            
            # 生成PDF
            doc.build(story)
            
            # 删除临时图像文件
            if os.path.exists(temp_img_path):
                os.remove(temp_img_path)
            
            messagebox.showinfo("成功", f"统计图表已导出到 {file_path}")
            
        except ImportError:
            messagebox.showerror("错误", "需要安装reportlab和Pillow库，请运行: pip install reportlab Pillow")
        except Exception as e:
            messagebox.showerror("错误", f"PDF导出失败: {str(e)}")

    def export_statistics_report(self):
        """导出统计报告"""
        # 询问保存位置
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            title="导出统计报告"
        )
        
        if not file_path:
            return
        
        # 统计数据
        total_count = len(self.data)
        subject_counts = {}
        difficulty_counts = {}
        
        for item in self.data:
            # 统计科目
            subject = item['科目']
            subject_counts[subject] = subject_counts.get(subject, 0) + 1
            
            # 统计难度
            difficulty = item['难度']
            difficulty_counts[difficulty] = difficulty_counts.get(difficulty, 0) + 1
        
        # 生成报告内容
        report = f"📊 错题统计报告\n"
        report += f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        report += f"总错题数: {total_count}\n\n"
        
        report += "各科目错题数:\n"
        for subject, count in sorted(subject_counts.items(), key=lambda x: x[1], reverse=True):
            percentage = (count / total_count * 100) if total_count > 0 else 0
            report += f"  {subject}: {count}题 ({percentage:.1f}%)\n"
        
        report += "\n各难度错题数:\n"
        for difficulty, count in sorted(difficulty_counts.items(), key=lambda x: x[1], reverse=True):
            percentage = (count / total_count * 100) if total_count > 0 else 0
            report += f"  {difficulty}: {count}题 ({percentage:.1f}%)\n"
        
        # 写入报告文件
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(report)
        
        messagebox.showinfo("成功", f"统计报告已导出到 {file_path}")
    
    def export_statistics_report_pdf(self):
        """导出统计报告为PDF"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import os
            
            # 注册中文字体
            font_name = 'SimSun'
            try:
                # 尝试注册字体
                pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))
            except:
                # 如果SimSun不可用，尝试其他常见中文字体
                try:
                    # 在Windows系统上查找常见字体文件
                    font_paths = [
                        r'C:\Windows\Fonts\simsun.ttc',    # 宋体
                        r'C:\Windows\Fonts\msyh.ttc',      # 微软雅黑
                        r'C:\Windows\Fonts\msyhbd.ttc',    # 微软雅黑粗体
                        r'C:\Windows\Fonts\simhei.ttf',    # 黑体
                    ]
                    
                    font_found = False
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            if 'simsun' in font_path.lower():
                                font_name = 'SimSun'
                            elif 'msyh' in font_path.lower():
                                font_name = 'MicrosoftYaHei'
                            elif 'simhei' in font_path.lower():
                                font_name = 'SimHei'
                            
                            pdfmetrics.registerFont(TTFont(font_name, font_path))
                            font_found = True
                            break
                    
                    if not font_found:
                        # 如果找不到系统字体，则报错
                        messagebox.showerror("错误", "未找到支持中文的字体文件")
                        return
                except:
                    messagebox.showerror("错误", "无法注册中文字体，请确保系统中有中文字体文件")
                    return
            
            # 询问保存位置
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
                title="导出统计报告为PDF"
            )
            
            if not file_path:
                return
            
            # 统计数据
            total_count = len(self.data)
            subject_counts = {}
            difficulty_counts = {}
            
            for item in self.data:
                # 统计科目
                subject = item['科目']
                subject_counts[subject] = subject_counts.get(subject, 0) + 1
                
                # 统计难度
                difficulty = item['难度']
                difficulty_counts[difficulty] = difficulty_counts.get(difficulty, 0) + 1
            
            # 创建PDF文档
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            story = []
            
            # 标题样式（使用中文字体）
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=getSampleStyleSheet()['Title'],
                fontName=font_name,
                fontSize=20,
                spaceAfter=30,
                alignment=1  # 居中
            )
            
            # 正文样式（使用中文字体）
            content_style = ParagraphStyle(
                'CustomContent',
                parent=getSampleStyleSheet()['Normal'],
                fontName=font_name,
                fontSize=12,
                spaceAfter=12
            )
            
            # 添加标题
            title = Paragraph("📊 错题统计报告", title_style)
            story.append(title)
            story.append(Spacer(1, 0.2 * inch))
            
            # 添加生成时间
            time_para = Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", content_style)
            story.append(time_para)
            
            # 添加总错题数
            count_para = Paragraph(f"总错题数: {total_count}", content_style)
            story.append(count_para)
            story.append(Spacer(1, 0.2 * inch))
            
            # 添加各科目错题数
            subject_header = Paragraph("各科目错题数:", content_style)
            story.append(subject_header)
            
            for subject, count in sorted(subject_counts.items(), key=lambda x: x[1], reverse=True):
                percentage = (count / total_count * 100) if total_count > 0 else 0
                subject_para = Paragraph(f"  {subject}: {count}题 ({percentage:.1f}%)", content_style)
                story.append(subject_para)
            
            story.append(Spacer(1, 0.2 * inch))
            
            # 添加各难度错题数
            difficulty_header = Paragraph("各难度错题数:", content_style)
            story.append(difficulty_header)
            
            for difficulty, count in sorted(difficulty_counts.items(), key=lambda x: x[1], reverse=True):
                percentage = (count / total_count * 100) if total_count > 0 else 0
                difficulty_para = Paragraph(f"  {difficulty}: {count}题 ({percentage:.1f}%)", content_style)
                story.append(difficulty_para)
            
            # 生成PDF
            doc.build(story)
            messagebox.showinfo("成功", f"统计报告已导出到 {file_path}")
            
        except ImportError:
            messagebox.showerror("错误", "需要安装reportlab库，请运行: pip install reportlab")
        except Exception as e:
            messagebox.showerror("错误", f"PDF导出失败: {str(e)}")
    
    def create_collection_manager_tab(self):
        """创建合集管理标签页"""
        # 创建框架
        self.collection_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.collection_frame, text="📚 合集管理")
        
        # 创建主框架
        main_frame = ttk.Frame(self.collection_frame)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建合集信息输入区域
        info_frame = ttk.LabelFrame(main_frame, text="➕ 创建新合集", padding=10)
        info_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 合集名称
        name_frame = ttk.Frame(info_frame)
        name_frame.pack(fill=tk.X, pady=5)
        ttk.Label(name_frame, text="合集名称:", font=('Microsoft YaHei', 10)).pack(side=tk.LEFT)
        self.collection_name_var = tk.StringVar()
        ttk.Entry(name_frame, textvariable=self.collection_name_var, font=('Microsoft YaHei', 10)).pack(side=tk.LEFT, padx=(10, 0), fill=tk.X, expand=True)
        
        # 合集描述
        desc_frame = ttk.Frame(info_frame)
        desc_frame.pack(fill=tk.X, pady=5)
        ttk.Label(desc_frame, text="合集描述:", font=('Microsoft YaHei', 10)).pack(side=tk.LEFT)
        self.collection_desc_var = tk.StringVar()
        ttk.Entry(desc_frame, textvariable=self.collection_desc_var, font=('Microsoft YaHei', 10)).pack(side=tk.LEFT, padx=(10, 0), fill=tk.X, expand=True)
        
        # 按钮区域
        button_frame = ttk.Frame(info_frame)
        button_frame.pack(fill=tk.X, pady=10)
        ttk.Button(button_frame, text="创建合集", command=self.create_collection, style='Action.TButton').pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(button_frame, text="重置", command=self.reset_collection_form, style='Action.TButton').pack(side=tk.LEFT)
        
        # 合集列表区域
        list_frame = ttk.LabelFrame(main_frame, text="📚 合集列表", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # 创建合集列表树形控件
        columns = ('名称', '描述', '题目数')
        self.collection_tree = ttk.Treeview(list_frame, columns=columns, show='headings', height=8)
        
        # 设置列标题和宽度
        col_widths = {'名称': 150, '描述': 300, '题目数': 100}
        for col in columns:
            self.collection_tree.heading(col, text=col, anchor=tk.CENTER)
            self.collection_tree.column(col, width=col_widths.get(col, 100), anchor=tk.W)
        
        # 添加滚动条
        v_scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.collection_tree.yview)
        self.collection_tree.configure(yscrollcommand=v_scrollbar.set)
        
        # 布局
        self.collection_tree.grid(row=0, column=0, sticky='nsew')
        v_scrollbar.grid(row=0, column=1, sticky='ns')
        
        list_frame.grid_rowconfigure(0, weight=1)
        list_frame.grid_columnconfigure(0, weight=1)
        
        # 操作按钮
        action_frame = ttk.Frame(list_frame)
        action_frame.grid(row=1, column=0, columnspan=2, pady=10, sticky='ew')
        
        # 为action_frame配置列权重
        action_frame.grid_columnconfigure(0, weight=1)
        action_frame.grid_columnconfigure(1, weight=1)
        action_frame.grid_columnconfigure(2, weight=1)
        
        ttk.Button(action_frame, text="查看题目", command=self.view_collection_problems, style='Action.TButton').grid(row=0, column=0, padx=2, sticky='ew')
        ttk.Button(action_frame, text="导出合集", command=self.export_collection, style='Action.TButton').grid(row=0, column=1, padx=2, sticky='ew')
        ttk.Button(action_frame, text="删除合集", command=self.delete_collection, style='Action.TButton').grid(row=0, column=2, padx=2, sticky='ew')
        
        # 加载现有合集
        self.refresh_collection_list()
    
    def create_collection(self):
        """创建新合集"""
        name = self.collection_name_var.get().strip()
        description = self.collection_desc_var.get().strip()
        
        if not name:
            messagebox.showwarning("警告", "请输入合集名称！")
            return
        
        # 检查合集是否已存在
        if name in self.collections:
            messagebox.showwarning("警告", "该合集名称已存在！")
            return
        
        # 创建新的合集
        self.collections[name] = {
            'description': description,
            'problems': []  # 存储题目ID或索引
        }
        
        # 保存数据
        self.save_data()
        
        # 刷新列表
        self.refresh_collection_list()
        
        # 重置表单
        self.reset_collection_form()
        
        messagebox.showinfo("成功", f"合集 '{name}' 创建成功！")
    
    def reset_collection_form(self):
        """重置合集表单"""
        self.collection_name_var.set("")
        self.collection_desc_var.set("")
    
    def refresh_collection_list(self):
        """刷新合集列表"""
        # 清空现有项目
        for item in self.collection_tree.get_children():
            self.collection_tree.delete(item)
        
        # 添加合集数据
        for name, info in self.collections.items():
            problem_count = len(info['problems'])
            self.collection_tree.insert('', tk.END, values=(name, info['description'], problem_count))
    
    def view_collection_problems(self):
        """查看合集中的题目"""
        selected_items = self.collection_tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择一个合集！")
            return
        
        # 获取选中的合集名称
        item = selected_items[0]
        values = self.collection_tree.item(item, 'values')
        collection_name = values[0]
        
        # 创建新窗口显示题目
        self.show_collection_problems_window(collection_name)
    
    def show_collection_problems_window(self, collection_name):
        """显示合集题目窗口"""
        # 创建新窗口
        window = tk.Toplevel(self.root)
        window.title(f"合集 '{collection_name}' 的题目")
        window.geometry("900x600")
        window.transient(self.root)
        window.grab_set()  # 模态窗口
        
        # 创建主框架
        main_frame = ttk.Frame(window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建表格
        columns = ('时间', '科目', '题干', '正确答案', '难度')
        tree = ttk.Treeview(main_frame, columns=columns, show='headings', height=15)
        
        # 设置列标题和宽度
        col_widths = {'时间': 140, '科目': 80, '题干': 250, '正确答案': 200, '难度': 80}
        for col in columns:
            tree.heading(col, text=col, anchor=tk.CENTER)
            tree.column(col, width=col_widths.get(col, 100), anchor=tk.W)
        
        # 添加滚动条
        v_scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=tree.yview)
        h_scrollbar = ttk.Scrollbar(main_frame, orient=tk.HORIZONTAL, command=tree.xview)
        tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        # 布局
        tree.grid(row=0, column=0, sticky='nsew')
        v_scrollbar.grid(row=0, column=1, sticky='ns')
        h_scrollbar.grid(row=1, column=0, sticky='ew')
        
        main_frame.grid_rowconfigure(0, weight=1)
        main_frame.grid_columnconfigure(0, weight=1)
        
        # 加载合集中的题目
        collection_info = self.collections.get(collection_name, {})
        problem_indices = collection_info.get('problems', [])
        
        for idx in problem_indices:
            if 0 <= idx < len(self.data):
                problem = self.data[idx]
                # 截断题干和答案以适应表格显示
                question = problem['题干'][:50] + '...' if len(problem['题干']) > 50 else problem['题干']
                answer = problem['正确答案'][:30] + '...' if len(problem['正确答案']) > 30 else problem['正确答案']
                
                tree.insert('', tk.END, values=(
                    problem['时间'],
                    problem['科目'],
                    question,
                    answer,
                    problem['难度']
                ))
    
    def delete_collection(self):
        """删除选中的合集"""
        selected_items = self.collection_tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择一个合集！")
            return
        
        # 获取选中的合集名称
        item = selected_items[0]
        values = self.collection_tree.item(item, 'values')
        collection_name = values[0]
        
        if messagebox.askyesno("确认", f"确定要删除合集 '{collection_name}' 吗？此操作不可恢复！"):
            # 从集合中删除合集
            if collection_name in self.collections:
                del self.collections[collection_name]
            
            # 从数据中移除合集标记
            for problem in self.data:
                if problem.get('合集') == collection_name:
                    problem['合集'] = ''
            
            # 保存数据
            self.save_data()
            
            # 刷新列表
            self.refresh_collection_list()
            
            messagebox.showinfo("成功", f"合集 '{collection_name}' 已删除！")
    
    def export_collection(self):
        """导出选中的合集"""
        selected_items = self.collection_tree.selection()
        if not selected_items:
            messagebox.showwarning("警告", "请先选择一个合集！")
            return
        
        # 获取选中的合集名称
        item = selected_items[0]
        values = self.collection_tree.item(item, 'values')
        collection_name = values[0]
        
        # 询问保存位置
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("Text files", "*.txt"), ("All files", "*.*")],
            title=f"导出合集 '{collection_name}'"
        )
        
        if not file_path:
            return
        
        # 获取合集中的题目
        collection_info = self.collections.get(collection_name, {})
        problem_indices = collection_info.get('problems', [])
        collection_problems = []
        
        for idx in problem_indices:
            if 0 <= idx < len(self.data):
                collection_problems.append(self.data[idx])
        
        if not collection_problems:
            messagebox.showwarning("警告", "该合集中没有题目！")
            return
        
        # 根据文件扩展名选择导出格式
        if file_path.lower().endswith('.csv'):
            self.export_collection_to_csv(collection_problems, file_path)
        elif file_path.lower().endswith('.txt'):
            self.export_collection_to_txt(collection_problems, file_path)
        else:
            # 默认使用CSV格式
            self.export_collection_to_csv(collection_problems, file_path)
    
    def export_collection_to_csv(self, problems, file_path):
        """将合集导出为CSV格式"""
        try:
            with open(file_path, 'w', newline='', encoding='utf-8-sig') as f:
                fieldnames = ['时间', '科目', '题干', '正确答案', '附件路径', '难度', '合集']
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(problems)
            
            messagebox.showinfo("成功", f"合集已导出到 {file_path}")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {str(e)}")
    
    def export_collection_to_txt(self, problems, file_path):
        """将合集导出为TXT格式"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"错题合集: {problems[0].get('合集', 'Unknown')}\n")
                f.write("="*50 + "\n\n")
                
                for i, problem in enumerate(problems, 1):
                    f.write(f"第{i}题\n")
                    f.write(f"时间: {problem['时间']}\n")
                    f.write(f"科目: {problem['科目']}\n")
                    f.write(f"难度: {problem['难度']}\n")
                    f.write(f"题干: {problem['题干']}\n")
                    f.write(f"正确答案: {problem['正确答案']}\n")
                    f.write(f"附件路径: {problem['附件路径'] if problem['附件路径'] else '无'}\n")
                    f.write("-" * 50 + "\n\n")
            
            messagebox.showinfo("成功", f"合集已导出到 {file_path}")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {str(e)}")
    
    def assign_problem_to_collection(self, problem_idx, collection_name):
        """将题目分配到指定合集"""
        if 0 <= problem_idx < len(self.data):
            # 更新题目的合集字段
            self.data[problem_idx]['合集'] = collection_name
            
            # 更新合集信息
            if collection_name not in self.collections:
                self.collections[collection_name] = {
                    'description': f'包含{collection_name}相关题目的合集',
                    'problems': []
                }
            
            # 添加题目索引到合集
            if problem_idx not in self.collections[collection_name]['problems']:
                self.collections[collection_name]['problems'].append(problem_idx)
    
    def remove_problem_from_collection(self, problem_idx, collection_name):
        """将题目从指定合集中移除"""
        if 0 <= problem_idx < len(self.data):
            # 更新题目的合集字段
            if self.data[problem_idx].get('合集') == collection_name:
                self.data[problem_idx]['合集'] = ''
            
            # 更新合集信息
            if collection_name in self.collections:
                if problem_idx in self.collections[collection_name]['problems']:
                    self.collections[collection_name]['problems'].remove(problem_idx)
                    
                    # 如果合集为空，可以考虑删除合集
                    if not self.collections[collection_name]['problems']:
                        # 不自动删除空合集，保留合集定义
                        pass

if __name__ == "__main__":
    root = tk.Tk()
    app = MistakeBookApp(root)
    root.mainloop()
