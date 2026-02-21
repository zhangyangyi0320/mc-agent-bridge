import csv
import random
from datetime import datetime, timedelta
import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import tkinter.font as tkFont

class TestDataGeneratorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("错题本测试数据生成器")
        self.root.geometry("600x800")
        
        # 初始化数据
        self.subjects = ["语文", "数学", "英语", "物理", "化学", "生物", "历史", "地理", "道法"]
        self.difficulties = ["简单", "中等", "困难"]
        
        # 默认值
        self.default_num_records = 100
        self.default_subject_ratios = {subject: 1/len(self.subjects) for subject in self.subjects}
        self.default_difficulty_ratios = {difficulty: 1/len(self.difficulties) for difficulty in self.difficulties}
        
        # 创建界面
        self.create_widgets()
        
    def create_widgets(self):
        """创建GUI界面"""
        # 主框架
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        title_label = ttk.Label(main_frame, text="错题本测试数据生成器", font=("Microsoft YaHei", 16, "bold"))
        title_label.pack(pady=(0, 20))
        
        # 条目数量设置
        num_frame = ttk.LabelFrame(main_frame, text="📊 数据条目数量", padding=10)
        num_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(num_frame, text="生成数量:").pack(side=tk.LEFT)
        self.num_var = tk.StringVar(value=str(self.default_num_records))
        num_spinbox = tk.Spinbox(num_frame, from_=1, to=10000, textvariable=self.num_var, width=10)
        num_spinbox.pack(side=tk.LEFT, padx=(10, 0))
        ttk.Label(num_frame, text="条").pack(side=tk.LEFT, padx=(5, 0))
        
        # 导出格式选择
        format_frame = ttk.LabelFrame(main_frame, text="💾 导出格式", padding=10)
        format_frame.pack(fill=tk.X, pady=5)
        
        self.export_formats = {
            'CSV': tk.BooleanVar(value=True),
            'TXT': tk.BooleanVar(value=True),
            'Excel': tk.BooleanVar(value=False),
            'Word': tk.BooleanVar(value=False)
        }
        
        for format_name, var in self.export_formats.items():
            ttk.Checkbutton(format_frame, text=format_name, variable=var).pack(side=tk.LEFT, padx=10)
        
        # 导出位置选择
        export_path_frame = ttk.LabelFrame(main_frame, text="📁 导出位置", padding=10)
        export_path_frame.pack(fill=tk.X, pady=5)
        
        self.export_path_var = tk.StringVar(value=os.getcwd())
        export_path_entry = ttk.Entry(export_path_frame, textvariable=self.export_path_var, width=50)
        export_path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        ttk.Button(export_path_frame, text="浏览", command=self.browse_export_path).pack(side=tk.LEFT)
        
        # 科目占比设置
        subject_frame = ttk.LabelFrame(main_frame, text="📚 科目占比设置", padding=10)
        subject_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 创建科目占比设置的滚动框架
        canvas = tk.Canvas(subject_frame, height=200)
        scrollbar = ttk.Scrollbar(subject_frame, orient="vertical", command=canvas.yview)
        subject_scrollable_frame = ttk.Frame(canvas)
        
        subject_scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=subject_scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 为每个科目创建比例输入框
        self.subject_ratio_vars = {}
        for i, subject in enumerate(self.subjects):
            row = i // 3  # 每行放3个
            col = (i % 3) * 2  # 每个控件占两列（标签+输入框）
            
            ttk.Label(subject_scrollable_frame, text=f"{subject}:").grid(row=row, column=col, sticky=tk.W, padx=5, pady=2)
            
            var = tk.StringVar(value="11.11")  # 默认约为1/9，即11.11%
            self.subject_ratio_vars[subject] = var
            
            entry = ttk.Entry(subject_scrollable_frame, textvariable=var, width=8)
            entry.grid(row=row, column=col+1, sticky=tk.W, padx=5, pady=2)
            ttk.Label(subject_scrollable_frame, text="%").grid(row=row, column=col+2, sticky=tk.W, padx=(0, 15), pady=2)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 难度占比设置
        difficulty_frame = ttk.LabelFrame(main_frame, text="📊 难度占比设置", padding=10)
        difficulty_frame.pack(fill=tk.X, pady=5)
        
        # 创建难度占比设置的框架
        self.difficulty_ratio_vars = {}
        for i, difficulty in enumerate(self.difficulties):
            ttk.Label(difficulty_frame, text=f"{difficulty}:").grid(row=0, column=i*2, sticky=tk.W, padx=5, pady=2)
            
            var = tk.StringVar(value="33.33")  # 默认约为1/3，即33.33%
            self.difficulty_ratio_vars[difficulty] = var
            
            entry = ttk.Entry(difficulty_frame, textvariable=var, width=8)
            entry.grid(row=0, column=i*2+1, sticky=tk.W, padx=5, pady=2)
            ttk.Label(difficulty_frame, text="%").grid(row=0, column=i*2+2, sticky=tk.W, padx=(0, 15), pady=2)
        
        # 生成按钮
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=20)
        
        ttk.Button(button_frame, text="🔄 重置为默认", command=self.reset_to_default).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(button_frame, text="⚙️ 自动平衡科目占比", command=self.auto_balance_ratios).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(button_frame, text="⚖️ 自动平衡难度占比", command=self.auto_balance_difficulty_ratios).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(button_frame, text="✨ 生成测试数据", command=self.generate_test_data, style='Accent.TButton').pack(side=tk.RIGHT)
    
    def browse_export_path(self):
        """浏览导出路径"""
        path = filedialog.askdirectory()
        if path:
            self.export_path_var.set(path)
    
    def reset_to_default(self):
        """重置为默认设置"""
        self.num_var.set(str(self.default_num_records))
        for var in self.export_formats.values():
            var.set(False)
        self.export_formats['CSV'].set(True)
        self.export_formats['TXT'].set(True)
        self.export_path_var.set(os.getcwd())
        
        # 重置科目比例为平均分布
        default_subject_ratio = round(100.0 / len(self.subjects), 2)
        for var in self.subject_ratio_vars.values():
            var.set(str(default_subject_ratio))
        
        # 重置难度比例为平均分布
        default_difficulty_ratio = round(100.0 / len(self.difficulties), 2)
        for var in self.difficulty_ratio_vars.values():
            var.set(str(default_difficulty_ratio))
    
    def auto_balance_ratios(self):
        """自动平衡科目占比为总和100%"""
        try:
            # 计算当前总和
            current_sum = sum(float(var.get()) for var in self.subject_ratio_vars.values())
            
            # 如果总和不为0，则按比例调整到总和为100
            if current_sum > 0:
                factor = 100.0 / current_sum
                for subject, var in self.subject_ratio_vars.items():
                    current_val = float(var.get())
                    new_val = round(current_val * factor, 2)
                    var.set(str(new_val))
            else:
                # 如果总和为0，设为平均分布
                default_ratio = round(100.0 / len(self.subjects), 2)
                for var in self.subject_ratio_vars.values():
                    var.set(str(default_ratio))
        except ValueError:
            messagebox.showerror("错误", "请输入有效的数值")
    
    def auto_balance_difficulty_ratios(self):
        """自动平衡难度占比为总和100%"""
        try:
            # 计算当前总和
            current_sum = sum(float(var.get()) for var in self.difficulty_ratio_vars.values())
            
            # 如果总和不为0，则按比例调整到总和为100
            if current_sum > 0:
                factor = 100.0 / current_sum
                for difficulty, var in self.difficulty_ratio_vars.items():
                    current_val = float(var.get())
                    new_val = round(current_val * factor, 2)
                    var.set(str(new_val))
            else:
                # 如果总和为0，设为平均分布
                default_ratio = round(100.0 / len(self.difficulties), 2)
                for var in self.difficulty_ratio_vars.values():
                    var.set(str(default_ratio))
        except ValueError:
            messagebox.showerror("错误", "请输入有效的数值")
    
    def generate_test_data(self):
        """生成测试数据"""
        try:
            # 获取参数
            num_records = int(self.num_var.get())
            export_path = self.export_path_var.get()
            
            # 获取科目比例
            subject_ratios = {}
            total_subject_ratio = 0
            for subject, var in self.subject_ratio_vars.items():
                ratio = float(var.get())
                subject_ratios[subject] = ratio
                total_subject_ratio += ratio
            
            # 验证科目比例总和
            if abs(total_subject_ratio - 100.0) > 0.1:  # 允许0.1的误差
                result = messagebox.askyesno(
                    "提示", 
                    f"科目占比总和为{total_subject_ratio:.2f}%，不等于100%。是否继续？\n\n" +
                    "点击'是'继续生成，点击'否'取消。"
                )
                if not result:
                    return
            
            # 获取难度比例
            difficulty_ratios = {}
            total_difficulty_ratio = 0
            for difficulty, var in self.difficulty_ratio_vars.items():
                ratio = float(var.get())
                difficulty_ratios[difficulty] = ratio
                total_difficulty_ratio += ratio
            
            # 验证难度比例总和
            if abs(total_difficulty_ratio - 100.0) > 0.1:  # 允许0.1的误差
                result = messagebox.askyesno(
                    "提示", 
                    f"难度占比总和为{total_difficulty_ratio:.2f}%，不等于100%。是否继续？\n\n" +
                    "点击'是'继续生成，点击'否'取消。"
                )
                if not result:
                    return
            
            # 生成数据
            test_data = self.generate_random_mistake_data(num_records, subject_ratios, difficulty_ratios)
            
            # 确定要导出的格式
            formats_to_export = [fmt for fmt, var in self.export_formats.items() if var.get()]
            
            if not formats_to_export:
                messagebox.showwarning("警告", "请至少选择一种导出格式！")
                return
            
            # 导出数据
            success_count = 0
            for fmt in formats_to_export:
                if fmt == 'CSV':
                    filename = os.path.join(export_path, "mistakebook_test_data.csv")
                    if self.save_data_to_csv(test_data, filename):
                        success_count += 1
                elif fmt == 'TXT':
                    filename = os.path.join(export_path, "mistakebook_test_data.txt")
                    if self.save_data_to_txt(test_data, filename):
                        success_count += 1
                elif fmt == 'Excel':
                    try:
                        import pandas as pd
                        filename = os.path.join(export_path, "mistakebook_test_data.xlsx")
                        if self.save_data_to_excel(test_data, filename):
                            success_count += 1
                    except ImportError:
                        messagebox.showerror("错误", "需要安装pandas和openpyxl库才能导出Excel格式：\npip install pandas openpyxl")
                elif fmt == 'Word':
                    try:
                        from docx import Document
                        from docx.shared import Inches
                        filename = os.path.join(export_path, "mistakebook_test_data.docx")
                        if self.save_data_to_word(test_data, filename):
                            success_count += 1
                    except ImportError:
                        messagebox.showerror("错误", "需要安装python-docx库才能导出Word格式：\npip install python-docx")
            
            # 显示完成信息
            if success_count > 0:
                messagebox.showinfo("成功", f"已成功生成 {len(test_data)} 条测试数据，\n导出到 {export_path} 目录下！")
            else:
                messagebox.showwarning("警告", "没有成功导出任何文件！")
                
        except ValueError as e:
            messagebox.showerror("错误", f"请输入有效的数值：{e}")
        except Exception as e:
            messagebox.showerror("错误", f"生成数据时出错：{e}")
    
    def generate_random_mistake_data(self, num_records, subject_ratios, difficulty_ratios=None):
        """
        生成与mistakebook兼容的随机测试数据
        
        Args:
            num_records (int): 要生成的记录数量
            subject_ratios (dict): 各科目的占比
            difficulty_ratios (dict): 各难度的占比
        
        Returns:
            list: 包含错题数据的字典列表
        """
        # 生成随机题干的词库
        question_banks = {
            "语文": [
                "下列句子中，加点字的读音完全正确的一项是",
                "下列词语中，没有错别字的一组是",
                "下列各句中，加点的成语使用恰当的一句是",
                "对下列句子中加点词语的解释，不正确的一项是",
                "下列各句中，没有语病的一句是",
                "下列文学常识的表述，错误的一项是",
                "古诗词鉴赏：阅读下面这首诗，完成题目",
                "文言文阅读：阅读下面的文言文，完成题目",
                "现代文阅读：阅读下面的文字，完成题目",
                "古诗文默写：补写出下列名篇名句中的空缺部分"
            ],
            "数学": [
                "已知函数 f(x) = x² + 2x + 1，则 f(x) 的最小值是",
                "若复数 z = (1+i)/(1-i)，则 |z| 的值为",
                "设等差数列 {an} 的前 n 项和为 Sn，若 a3 = 5, S5 = 25，则 a10 = ",
                "在三角形 ABC 中，若 sin A : sin B : sin C = 3 : 4 : 5，则 cos C = ",
                "设函数 f(x) = ln x - ax (a > 0)，则 f(x) 的单调递增区间为",
                "已知向量 a = (1, 2), b = (x, 4)，若 a ⊥ b，则 x = ",
                "椭圆 x²/4 + y²/3 = 1 的离心率为",
                "设随机变量 X ~ N(2, 4)，则 P(X > 2) = ",
                "若 x, y 满足约束条件，则 z = x + y 的最大值为",
                "函数 f(x) = sin x + cos x 的最大值为"
            ],
            "英语": [
                "Choose the word that is closest in meaning to the underlined word.",
                "Choose the correct form of the verb to complete the sentence.",
                "Read the passage and answer the following questions.",
                "Choose the best answer to complete the dialogue.",
                "Identify the grammatical error in the following sentence.",
                "Choose the correct preposition to complete the sentence.",
                "Fill in the blank with the appropriate word.",
                "Choose the correct translation of the following sentence.",
                "Complete the sentence with the correct form of the word given in brackets.",
                "Select the best option to complete the sentence."
            ],
            "物理": [
                "一个物体做匀加速直线运动，初速度为2m/s，加速度为3m/s²，则5s后的速度为",
                "质量为2kg的物体在水平面上受到8N的水平拉力，若摩擦系数为0.2，则物体的加速度为",
                "一个理想变压器的原副线圈匝数比为10:1，若原线圈电压为220V，则副线圈电压为",
                "在单缝衍射实验中，当缝宽减小时，衍射现象会",
                "一个电子在匀强磁场中做匀速圆周运动，若磁感应强度增大一倍，则周期变为原来的",
                "理想气体在等温过程中，当体积减小时，压强会",
                "光电效应中，光电子的最大初动能与入射光的频率成",
                "在LC振荡电路中，当电容器的电荷量最大时，磁场能为",
                "一束光从空气射入水中，若入射角为30°，水的折射率为4/3，则折射角约为",
                "在双缝干涉实验中，若增大双缝间距，则条纹间距会"
            ],
            "化学": [
                "下列物质中，属于电解质的是",
                "在标准状况下，1mol任何气体的体积都约为",
                "下列反应中，属于氧化还原反应的是",
                "对于可逆反应 A(g) + B(g) ⇌ 2C(g)，增大压强平衡将",
                "下列离子在水溶液中能够大量共存的是",
                "酸碱中和滴定中，常用的指示剂是",
                "已知反应热化学方程式，求反应的焓变",
                "下列原子中，第一电离能最大的是",
                "在Fe(OH)₃胶体中逐滴加入某种溶液，先产生沉淀后沉淀溶解，则该溶液是",
                "下列有机物中，能发生银镜反应的是"
            ],
            "生物": [
                "细胞膜的主要成分是",
                "在有丝分裂过程中，染色体数目加倍发生在",
                "下列关于DNA复制的叙述，正确的是",
                "孟德尔遗传定律适用于",
                "人体的特异性免疫包括",
                "光合作用中，产生氧气的阶段是",
                "在生态系统中，碳元素主要以什么形式循环",
                "基因突变的特点不包括",
                "兴奋在神经纤维上的传导形式是",
                "生长素在植物体内的运输特点是"
            ],
            "历史": [
                "下列哪一事件标志着中国近代史的开端",
                "辛亥革命最重要的历史功绩是",
                "五四运动的导火索是",
                "中国共产党成立的标志是",
                "遵义会议的重要意义在于",
                "抗日战争全面爆发的标志是",
                "新中国成立的历史意义不包括",
                "改革开放政策是在哪次会议上提出的",
                "世界古代史上，亚历山大帝国的建立者是",
                "文艺复兴运动首先兴起于哪个国家"
            ],
            "地理": [
                "地球自转的方向是",
                "下列哪种地貌是由流水侵蚀作用形成的",
                "全球气候变暖的主要原因是",
                "下列哪个因素对气候的影响最小",
                "板块构造学说认为，地震多发生在",
                "下列哪个城市位于长江三角洲地区",
                "我国水资源空间分布的特点是",
                "影响城市区位的自然因素不包括",
                "世界最大的沙漠是",
                "我国四大盆地中，海拔最高的是"
            ],
            "道法": [
                "下列属于我国基本政治制度的是",
                "公民行使监督权的渠道不包括",
                "中国特色社会主义最本质的特征是",
                "我国的根本制度是",
                "社会主义核心价值观中，属于国家层面的是",
                "公民的基本义务包括",
                "人民法院是我国的",
                "我国的国家政权组织形式是",
                "民法典中新增的独立成编的是",
                "全面推进依法治国的总目标是"
            ]
        }
        
        # 生成随机答案的词库
        answer_banks = {
            "语文": [
                "根据汉字的读音规则和声调，正确答案是B选项。",
                "通过分析各选项中的汉字结构和用法，确定A项为正确答案。",
                "结合语境和成语的含义，C项使用恰当。",
                "根据上下文和词义辨析，选择最合适的解释。",
                "分析句子结构，找出语病类型并修正。",
                "熟悉文学常识，准确判断各选项的正误。",
                "从诗歌的意境、表达技巧等角度分析，得出答案。",
                "理解文言文中的实词、虚词、句式，推断文意。",
                "把握文章主旨，理解作者观点态度。",
                "熟练背诵名篇名句，准确填写。"
            ],
            "数学": [
                "利用配方法或求导法找到函数的最小值点。",
                "利用复数除法法则和模的定义计算。",
                "运用等差数列的通项公式和前n项和公式求解。",
                "利用正弦定理和余弦定理解决三角形问题。",
                "求导后分析导函数的符号确定单调区间。",
                "利用向量垂直的条件 a·b = 0 求解。",
                "根据椭圆的标准方程和离心率公式计算。",
                "利用正态分布的对称性求解概率。",
                "利用线性规划的方法求目标函数的最值。",
                "将函数化为辅助角公式的形式求最值。"
            ],
            "英语": [
                "根据词汇的含义和语境选择同义词。",
                "根据语法规则选择正确的动词形式。",
                "仔细阅读文章，理解段落大意和细节信息。",
                "根据对话的语境和习惯用法选择最佳答案。",
                "识别句子中的语法错误类型。",
                "根据句意选择合适的介词。",
                "根据上下文选择合适的词汇。",
                "根据语法规则和词汇含义选择正确翻译。",
                "根据语法规则变化单词形式。",
                "根据语法规则和语境选择最佳选项。"
            ],
            "物理": [
                "使用匀变速直线运动的速度公式 v = v₀ + at 计算。",
                "根据牛顿第二定律 F_net = ma 计算加速度。",
                "利用变压器的电压比等于匝数比的原理计算。",
                "根据单缝衍射规律，缝宽越小衍射越明显。",
                "根据洛伦兹力提供向心力，T = 2πm/(qB)。",
                "根据玻意耳定律，等温过程 pV = 常数。",
                "根据爱因斯坦光电效应方程分析。",
                "在LC振荡电路中，电场能和磁场能相互转化。",
                "根据折射定律 n₁sinθ₁ = n₂sinθ₂ 计算。",
                "根据双缝干涉条纹间距公式 Δx = λL/d 分析。"
            ],
            "化学": [
                "电解质是在水溶液中或熔融状态下能导电的化合物。",
                "标准状况下(0°C, 101kPa)，1mol气体体积约为22.4L。",
                "氧化还原反应中元素化合价发生变化。",
                "根据勒夏特列原理，增大压强平衡向气体分子数减少的方向移动。",
                "考虑离子间是否发生反应生成沉淀、气体或弱电解质。",
                "常用的酸碱指示剂包括酚酞、甲基橙等。",
                "根据盖斯定律，利用已知热化学方程式计算。",
                "同周期元素从左到右第一电离能逐渐增大。",
                "先加入电解质聚沉胶体，再加入酸溶解沉淀。",
                "含有醛基的有机物能发生银镜反应。"
            ],
            "生物": [
                "细胞膜主要由磷脂双分子层和蛋白质构成。",
                "在有丝分裂后期，着丝点分裂，染色体数目加倍。",
                "DNA复制是半保留复制，需要解旋酶和DNA聚合酶。",
                "孟德尔遗传定律适用于进行有性生殖的真核生物核基因遗传。",
                "人体特异性免疫包括体液免疫和细胞免疫。",
                "在光合作用的光反应阶段产生氧气。",
                "碳元素在生物圈中以CO₂的形式循环。",
                "基因突变具有随机性、不定向性、低频性等特点。",
                "兴奋在神经纤维上以局部电流的形式传导。",
                "生长素在植物体内进行极性运输。"
            ],
            "历史": [
                "1840年鸦片战争爆发标志着中国近代史的开端。",
                "辛亥革命推翻了清朝统治，结束了中国两千多年的封建帝制。",
                "巴黎和会上中国外交失败成为五四运动的导火索。",
                "1921年中共一大的召开标志着中国共产党的成立。",
                "遵义会议确立了毛泽东在党内的领导地位。",
                "1937年卢沟桥事变标志着抗日战争全面爆发。",
                "新中国成立结束了中国半殖民地半封建社会的历史。",
                "1978年十一届三中全会提出改革开放政策。",
                "亚历山大帝国由亚历山大大帝建立。",
                "文艺复兴运动首先兴起于意大利。"
            ],
            "地理": [
                "地球自西向东自转，从北极上空看呈逆时针方向。",
                "流水侵蚀作用形成峡谷、瀑布等地貌。",
                "大量燃烧化石燃料导致温室气体增加。",
                "太阳辐射、大气环流、地面状况等影响气候，洋流影响相对较小。",
                "地震多发生在板块交界处。",
                "上海位于长江三角洲地区。",
                "我国水资源南多北少，东多西少。",
                "自然因素包括地形、气候、河流等，交通属于社会经济因素。",
                "撒哈拉沙漠是世界最大的沙漠。",
                "柴达木盆地是我国海拔最高的盆地。"
            ],
            "道法": [
                "我国的基本政治制度包括人民代表大会制度、中国共产党领导的多党合作和政治协商制度等。",
                "公民行使监督权的渠道包括人大代表联系群众制度、信访举报制度等。",
                "中国特色社会主义最本质的特征是中国共产党领导。",
                "我国的根本制度是社会主义制度。",
                "富强、民主、文明、和谐是国家层面的价值目标。",
                "公民的基本义务包括维护国家统一和民族团结等。",
                "人民法院是我国的审判机关。",
                "我国的政权组织形式是人民代表大会制度。",
                "民法典中人格权独立成编是重大创新。",
                "全面推进依法治国的总目标是建设中国特色社会主义法治体系。"
            ]
        }
        
        data = []
        
        # 根据比例生成科目列表
        subject_list = []
        for subject, ratio in subject_ratios.items():
            count = int(num_records * ratio / 100)
            subject_list.extend([subject] * count)
        
        # 如果数量不够，用随机科目补齐
        while len(subject_list) < num_records:
            subject_list.append(random.choice(self.subjects))
        
        # 打乱科目顺序
        random.shuffle(subject_list)
        
        # 根据比例生成难度列表
        difficulty_list = []
        if difficulty_ratios:  # 如果提供了难度比例
            for difficulty, ratio in difficulty_ratios.items():
                count = int(num_records * ratio / 100)
                difficulty_list.extend([difficulty] * count)
            
            # 如果数量不够，用随机难度补齐
            while len(difficulty_list) < num_records:
                difficulty_list.append(random.choice(self.difficulties))
            
            # 打乱难度顺序
            random.shuffle(difficulty_list)
        else:
            # 如果没有提供难度比例，则随机分配
            difficulty_list = [random.choice(self.difficulties) for _ in range(num_records)]
        
        for i in range(num_records):
            # 按照预设比例选择科目
            subject = subject_list[i]
            
            # 按照预设比例选择难度
            difficulty = difficulty_list[i]
            
            # 随机生成时间（最近30天内）
            start_date = datetime.now() - timedelta(days=30)
            random_date = start_date + timedelta(days=random.randint(0, 30), 
                                               hours=random.randint(0, 23), 
                                               minutes=random.randint(0, 59),
                                               seconds=random.randint(0, 59))
            time_str = random_date.strftime("%Y-%m-%d %H:%M:%S")
            
            # 随机选择题干和答案
            question = random.choice(question_banks[subject])
            answer = random.choice(answer_banks[subject])
            
            # 附件路径（随机生成或为空）
            attachment_path = f"attachment_{i+1}.pdf" if random.random() > 0.7 else ""
            
            # 创建数据记录
            record = {
                '时间': time_str,
                '科目': subject,
                '题干': question,
                '正确答案': answer,
                '附件路径': attachment_path,
                '难度': difficulty
            }
            
            data.append(record)
        
        return data

    def save_data_to_csv(self, data, filename):
        """
        将数据保存为CSV文件
        
        Args:
            data (list): 要保存的数据列表
            filename (str): 保存的文件名
        """
        try:
            fieldnames = ['时间', '科目', '题干', '正确答案', '附件路径', '难度']
            
            with open(filename, 'w', newline='', encoding='utf-8-sig') as csvfile:
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(data)
            
            print(f"成功生成 {len(data)} 条测试数据，保存到 {filename}")
            return True
        except Exception as e:
            print(f"保存CSV文件失败: {e}")
            return False

    def save_data_to_txt(self, data, filename):
        """
        将数据保存为TXT文件（用于测试导入功能）
        
        Args:
            data (list): 要保存的数据列表
            filename (str): 保存的文件名
        """
        try:
            with open(filename, 'w', encoding='utf-8') as txtfile:
                txtfile.write("错题本测试数据\n")
                txtfile.write("="*50 + "\n\n")
                
                for i, row in enumerate(data, 1):
                    txtfile.write(f"第{i}题\n")
                    txtfile.write(f"时间: {row['时间']}\n")
                    txtfile.write(f"科目: {row['科目']}\n")
                    txtfile.write(f"难度: {row['难度']}\n")
                    txtfile.write(f"题干: {row['题干']}\n")
                    txtfile.write(f"正确答案: {row['正确答案']}\n")
                    txtfile.write(f"附件路径: {row['附件路径'] if row['附件路径'] else '无'}\n")
                    txtfile.write("-" * 50 + "\n\n")
            
            print(f"成功生成 {len(data)} 条测试数据，保存到 {filename}")
            return True
        except Exception as e:
            print(f"保存TXT文件失败: {e}")
            return False

    def save_data_to_excel(self, data, filename):
        """
        将数据保存为Excel文件
        
        Args:
            data (list): 要保存的数据列表
            filename (str): 保存的文件名
        """
        try:
            import pandas as pd
            df = pd.DataFrame(data)
            df.to_excel(filename, index=False)
            print(f"成功生成 {len(data)} 条测试数据，保存到 {filename}")
            return True
        except Exception as e:
            print(f"保存Excel文件失败: {e}")
            return False

    def save_data_to_word(self, data, filename):
        """
        将数据保存为Word文档
        
        Args:
            data (list): 要保存的数据列表
            filename (str): 保存的文件名
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 创建文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('错题本测试数据', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加所有数据
            for i, row in enumerate(data, 1):
                # 添加题号标题
                doc.add_heading(f'第{i}题', level=1)
                
                # 添加各项内容
                doc.add_paragraph(f"时间: {row['时间']}")
                doc.add_paragraph(f"科目: {row['科目']}")
                doc.add_paragraph(f"难度: {row['难度']}")
                
                # 题干可能较长，单独段落
                p_question = doc.add_paragraph()
                p_question.add_run(f"题干: ").bold = True
                p_question.add_run(f"{row['题干']}")
                
                # 正确答案
                p_answer = doc.add_paragraph()
                p_answer.add_run(f"正确答案: ").bold = True
                p_answer.add_run(f"{row['正确答案']}")
                
                # 附件路径
                doc.add_paragraph(f"附件路径: {row['附件路径'] if row['附件路径'] else '无'}")
                
                # 添加分隔线
                doc.add_paragraph("-" * 50)
                doc.add_paragraph()  # 空行分隔
            
            # 保存文档
            doc.save(filename)
            print(f"成功生成 {len(data)} 条测试数据，保存到 {filename}")
            return True
        except Exception as e:
            print(f"保存Word文件失败: {e}")
            return False

    def save_data_to_pdf(self, data, filename):
        """
        将数据保存为PDF文件
        
        Args:
            data (list): 要保存的数据列表
            filename (str): 保存的文件名
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import os
            
            # 注册中文字体
            font_name = 'SimSun'
            try:
                # 尝试注册字体
                pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))
            except:
                # 如果SimSun不可用，尝试其他常见中文字体
                try:
                    # 在Windows系统上查找常见字体文件
                    font_paths = [
                        r'C:\Windows\Fonts\simsun.ttc',    # 宋体
                        r'C:\Windows\Fonts\msyh.ttc',      # 微软雅黑
                        r'C:\Windows\Fonts\msyhbd.ttc',    # 微软雅黑粗体
                        r'C:\Windows\Fonts\simhei.ttf',    # 黑体
                    ]
                    
                    font_found = False
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            if 'simsun' in font_path.lower():
                                font_name = 'SimSun'
                            elif 'msyh' in font_path.lower():
                                font_name = 'MicrosoftYaHei'
                            elif 'simhei' in font_path.lower():
                                font_name = 'SimHei'
                            
                            pdfmetrics.registerFont(TTFont(font_name, font_path))
                            font_found = True
                            break
                    
                    if not font_found:
                        # 如果找不到系统字体，则报错
                        print("未找到支持中文的字体文件")
                        return False
                except:
                    print("无法注册中文字体，请确保系统中有中文字体文件")
                    return False
            
            doc = SimpleDocTemplate(filename, pagesize=A4)
            story = []
            
            # 标题样式（使用中文字体）
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=getSampleStyleSheet()['Title'],
                fontName=font_name,
                fontSize=20,
                spaceAfter=30,
                alignment=1  # 居中
            )
            
            # 正文样式（使用中文字体）
            content_style = ParagraphStyle(
                'CustomContent',
                parent=getSampleStyleSheet()['Normal'],
                fontName=font_name,
                fontSize=10,
                spaceAfter=6
            )
            
            # 添加标题
            title = Paragraph("错题本测试数据", title_style)
            story.append(title)
            story.append(Spacer(1, 0.2 * inch))
            
            # 添加数据
            for i, row in enumerate(data, 1):
                story.append(Paragraph(f"第{i}题", content_style))
                story.append(Paragraph(f"时间: {row['时间']}", content_style))
                story.append(Paragraph(f"科目: {row['科目']}", content_style))
                story.append(Paragraph(f"难度: {row['难度']}", content_style))
                story.append(Paragraph(f"题干: {row['题干']}", content_style))
                story.append(Paragraph(f"正确答案: {row['正确答案']}", content_style))
                story.append(Paragraph(f"附件路径: {row['附件路径'] if row['附件路径'] else '无'}", content_style))
                story.append(Spacer(1, 0.2 * inch))
            
            # 生成PDF
            doc.build(story)
            print(f"成功生成 {len(data)} 条测试数据，保存到 {filename}")
            return True
        except ImportError:
            print("需要安装reportlab库，请运行: pip install reportlab")
            return False
        except Exception as e:
            print(f"保存PDF文件失败: {e}")
            return False

def main():
    root = tk.Tk()
    app = TestDataGeneratorGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()